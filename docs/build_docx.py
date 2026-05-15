#!/usr/bin/env python3
"""Convert the PRISM & Job Blueprint Process Guide markdown to a Word document."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).parent
MD_PATH = SCRIPT_DIR / "PRISM_and_Job_Blueprint_Process_Guide.md"
OUT_PATH = SCRIPT_DIR / "PRISM_and_Job_Blueprint_Process_Guide.docx"

# ── Styles ──────────────────────────────────────────────────────────

def setup_styles(doc: Document):
    """Configure document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        heading = doc.styles[f'Heading {level}']
        heading.font.name = 'Calibri'
        heading.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            heading.font.size = Pt(24)
            heading.paragraph_format.space_before = Pt(24)
            heading.paragraph_format.space_after = Pt(12)
        elif level == 2:
            heading.font.size = Pt(18)
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(8)
        elif level == 3:
            heading.font.size = Pt(14)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
        else:
            heading.font.size = Pt(12)
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(4)

    # Code style
    if 'Code' not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style('Code', 1)  # paragraph style
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(0xd4, 0xd4, 0xd4)
        code_style.paragraph_format.space_before = Pt(2)
        code_style.paragraph_format.space_after = Pt(2)


def set_cell_shading(cell, color_hex: str):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_table_from_lines(doc: Document, header_line: str, rows: list[str]):
    """Parse markdown table lines into a Word table."""
    def parse_cells(line: str) -> list[str]:
        return [c.strip() for c in line.strip('|').split('|')]

    headers = parse_cells(header_line)
    num_cols = len(headers)

    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1a1a2e')

    # Data rows
    for row_line in rows:
        cells_data = parse_cells(row_line)
        # Pad or trim to match header count
        while len(cells_data) < num_cols:
            cells_data.append('')
        cells_data = cells_data[:num_cols]

        row = table.add_row()
        for i, val in enumerate(cells_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold markers
            val_clean = val.replace('**', '')
            run = p.add_run(val_clean)
            run.font.size = Pt(10)
            if '**' in val:
                run.bold = True

    return table


def add_code_block(doc: Document, code_lines: list[str]):
    """Add a formatted code block."""
    code_text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x1e)

    # Add shading to paragraph
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'f5f5f5',
        qn('w:val'): 'clear',
    })
    pPr.append(shd)


def add_rich_paragraph(doc: Document, text: str, is_bullet: bool = False, bullet_level: int = 0):
    """Add a paragraph with bold/code inline formatting."""
    if is_bullet:
        p = doc.add_paragraph(style='List Bullet')
        if bullet_level > 0:
            p.style = doc.styles['List Bullet 2'] if 'List Bullet 2' in [s.name for s in doc.styles] else doc.styles['List Bullet']
    else:
        p = doc.add_paragraph()

    # Split on bold markers and backticks
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        else:
            p.add_run(part)

    return p


# ── Main Parser ─────────────────────────────────────────────────────

def convert_md_to_docx(md_path: Path, out_path: Path):
    doc = Document()
    setup_styles(doc)

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    lines = md_path.read_text(encoding='utf-8').split('\n')

    i = 0
    in_code_block = False
    code_lines: list[str] = []
    table_header: str | None = None
    table_rows: list[str] = []

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ──
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if table_header:
                    add_table_from_lines(doc, table_header, table_rows)
                    table_header = None
                    table_rows = []
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Tables ──
        if '|' in line and line.strip().startswith('|'):
            stripped = line.strip()
            # Skip separator lines
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue

            if table_header is None:
                table_header = stripped
            else:
                table_rows.append(stripped)
            i += 1
            continue
        else:
            # Flush pending table
            if table_header:
                add_table_from_lines(doc, table_header, table_rows)
                table_header = None
                table_rows = []

        # ── Headings ──
        if line.startswith('# ') and not line.startswith('##'):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # ── Horizontal rules ──
        if line.strip() == '---':
            doc.add_paragraph()  # spacing
            i += 1
            continue

        # ── Bullet points ──
        stripped = line.strip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            level = 0
            if line.startswith('   ') or line.startswith('\t'):
                level = 1
            add_rich_paragraph(doc, stripped[2:], is_bullet=True, bullet_level=level)
            i += 1
            continue

        # ── Numbered lists ──
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', num_match.group(2))
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
                else:
                    p.add_run(part)
            i += 1
            continue

        # ── Empty lines ──
        if stripped == '':
            i += 1
            continue

        # ── Regular paragraphs ──
        add_rich_paragraph(doc, stripped)
        i += 1

    # Flush any remaining table
    if table_header:
        add_table_from_lines(doc, table_header, table_rows)

    # ── Title page (insert at beginning) ──
    # We'll add metadata to the first paragraph area instead
    # Save
    doc.save(str(out_path))
    print(f"Created: {out_path}")
    print(f"Size: {out_path.stat().st_size / 1024:.0f} KB")


if __name__ == '__main__':
    convert_md_to_docx(MD_PATH, OUT_PATH)
