"""Generate New UI Adoption Plan Word document for Inspire Genius."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0, 51, 102)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def bold_para(text, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Inspire Genius\nNew UI Adoption Plan")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    "Role-Based Single-Page PWA Dashboard\n"
    "User \u2022 Manager \u2022 Company Admin \u2022 Practitioner \u2022 Distributor"
)
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0, 128, 128)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Prepared for: ").bold = True
meta.add_run("AES / Inspire Genius Team\n")
meta.add_run("Date: ").bold = True
meta.add_run("March 14, 2026\n")
meta.add_run("Classification: ").bold = True
meta.add_run("Internal \u2014 Implementation Planning")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Current UI Architecture",
    "3. New UI Architecture Overview",
    "4. Role Definitions & Multi-Role Logic",
    "5. Section 1: User Dashboard \u2014 Detailed Breakdown",
    "6. Section 2: Manager Dashboard \u2014 Detailed Breakdown",
    "7. Section 3: Company Admin Dashboard \u2014 Detailed Breakdown",
    "8. Section 4: Practitioner Dashboard \u2014 Detailed Breakdown",
    "9. Section 5: Distributor Dashboard \u2014 Detailed Breakdown",
    "10. Section 6: Super Admin (Future)",
    "11. Existing Functionality Mapping",
    "12. New Features Required",
    "13. Components to Modify",
    "14. New Components to Build",
    "15. Routing & Navigation Changes",
    "16. Data Layer: New Services, Hooks & Types",
    "17. Cost Dashboard Integration",
    "18. Implementation Phases",
    "19. Risk Assessment",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Number")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "Inspire Genius is transitioning from its current dual-layout architecture (UserLayout + "
    "SuperAdminLayout with 2 primary roles) to a unified single-page PWA with 5 role-based "
    "dashboard sections: User, Manager, Company Admin, Practitioner, and Distributor. A 6th "
    "section for Super Admin will be added later."
)
doc.add_paragraph(
    "The new UI introduces significant new functionality: PRISM behavioral assessments with "
    "team dynamics, practitioner-client coaching workflows, credit-based licensing for distributors, "
    "manager hiring/interview pipelines, company-wide org charts and training programs, and a "
    "cost analytics dashboard. Users who hold multiple roles will see all corresponding sections "
    "rendered in a unified navigation sidebar."
)
doc.add_paragraph(
    "This document provides a detailed mapping of every component in the new UI prototype "
    "(hosted at ig-interactive-dashboard-dev.s3-website-us-east-1.amazonaws.com) to the existing "
    "codebase, identifies what can be reused, what needs modification, and what must be built from scratch."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. CURRENT UI ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Current UI Architecture", level=1)

doc.add_heading("2.1 Current Roles", level=2)
add_table(
    ["Role", "Constant", "Layout", "Navigation Items"],
    [
        ["User", 'ROLES.USER = "user"', "UserLayout", "Home, Chat with Coaches, Manage Coaches, Documents, Settings, Help"],
        ["Super Admin", 'ROLES.SUPER_ADMIN = "super-admin"', "SuperAdminLayout", "Dashboard, User Mgmt, Coach Mgmt, RLHF Training, Prompt Builder, Audit Log, Settings, Project Log"],
        ["Admin", 'ROLES.ADMIN = "admin"', "SuperAdminLayout", "Same as Super Admin"],
        ["Company Admin", 'ROLES.COMPANY_ADMIN = "company-admin"', "SuperAdminLayout (subset)", "Dashboard, Users, Coaches, Settings"],
        ["Manager Admin", 'ROLES.MANAGER_ADMIN = "manager-admin"', "SuperAdminLayout (subset)", "Dashboard, Users"],
    ],
)

doc.add_heading("2.2 Current Pages", level=2)
add_table(
    ["Page", "Route", "Description"],
    [
        ["Home", "/home", "Welcome banner, stat cards, explore coaches, PRISM video, Ask Alex"],
        ["Dashboard (Coach List)", "/dashboard", "Coach chat interface with conversation list"],
        ["Coach Chat", "/dashboard/:coach/chat", "Full chat window with voice, documents, export"],
        ["Coaches", "/coaches", "Manage assigned coaches, preferences (tone, accent, gender)"],
        ["Documents", "/documents", "Upload, view, delete documents"],
        ["Settings", "/settings", "User profile settings"],
        ["Help", "/help", "Support tickets / issue reporting"],
        ["SA Dashboard", "/super-admin/dashboard", "System stats, org/user/license overview"],
        ["SA User Mgmt", "/super-admin/users", "Invite, edit, delete users, role assignment"],
        ["SA Coach Mgmt", "/super-admin/coaches", "Create, edit, deactivate coaches"],
        ["SA RLHF Training", "/super-admin/rlhf-training", "Feedback metrics, rating charts, feedback table"],
        ["SA Prompt Builder", "/super-admin/prompt-builder", "System prompt wizard with live preview"],
        ["SA Audit Log", "/super-admin/audit-log", "Event logging viewer with filters"],
        ["SA Organizations", "/super-admin/organizations", "Org CRUD, coach assignment (commented out)"],
        ["SA Settings", "/super-admin/settings", "Admin settings"],
    ],
)

doc.add_heading("2.3 Current Navigation Structure", level=2)
doc.add_paragraph(
    "Two completely separate layouts: UserLayout (sidebar with 6 items) and SuperAdminLayout "
    "(sidebar with 8 items). Users see one or the other based on their role. The new UI replaces "
    "this with a single unified sidebar that shows sections based on all roles a user holds."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. NEW UI ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. New UI Architecture Overview", level=1)

doc.add_heading("3.1 Single-Page PWA Shell", level=2)
doc.add_paragraph(
    "The new interface is a single-page application with a persistent sidebar navigation. "
    "The sidebar is organized into titled groups (MAIN, MANAGER, COMPANY, PRACTITIONER, "
    "DISTRIBUTOR, TOOLS, ACCOUNT). Each group corresponds to a role-based section. "
    "The sidebar, header, and content area persist across all views."
)

doc.add_heading("3.2 Layout Structure", level=2)
bullet("Sidebar (w-48, fixed left): Logo, role-grouped navigation items, user profile avatar at bottom")
bullet("Header (sticky top): Breadcrumb trail, date range filter dropdown, Export button")
bullet("Main Content (flex-1, scrollable): Dashboard content that changes based on active route")
bullet("Color scheme: Teal (#00A3A3) primary accent, white cards, gray-50 background, Inter font")

doc.add_heading("3.3 Navigation Groups (from new UI prototype)", level=2)
add_table(
    ["Group Title", "Role Required", "Navigation Items", "Routes"],
    [
        ["MAIN", "User (all users)", "Dashboard, Chat, PRISM Report, Documents, Goals", "/dashboard, /chat, /prism-report, /documents, /goals"],
        ["MANAGER", "Manager", "Manager Dashboard, My Team, Career Mgmt, Hiring, Team Building, Training, Leadership, Interviews, Candidates, Job DNA", "/manager/dashboard, /manager/team, /manager/career, /manager/hiring, /manager/team-building, /manager/training, /manager/leadership, /manager/interviews, /manager/candidates, /manager/job-dna"],
        ["COMPANY", "Company Admin", "Company Dashboard, Org Chart, Teams, Training, Assignments, Leadership Profiles", "/company/dashboard, /company/org-chart, /company/teams, /company/training, /company/assignments, /company/leadership"],
        ["PRACTITIONER", "Practitioner", "Practitioner Home, My Clients, Conversations, Follow-Up, Credits", "/practitioner/dashboard, /practitioner/clients, /practitioner/conversations, /practitioner/follow-up, /practitioner/credits"],
        ["DISTRIBUTOR", "Distributor", "Distributor Home, Practitioners, Credits, Allocate Credits", "/distributor/dashboard, /distributor/practitioners, /distributor/credits, /distributor/credit-allocation"],
        ["TOOLS", "All users", "Assessment, Settings", "/assessment, /settings"],
        ["ACCOUNT", "All users", "Notifications, Messages", "/notifications, /messages"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. ROLE DEFINITIONS & MULTI-ROLE LOGIC
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Role Definitions & Multi-Role Logic", level=1)

doc.add_heading("4.1 Expanded Role Definitions", level=2)
add_table(
    ["Role", "Constant", "Description", "New UI Sections Visible"],
    [
        ["User", '"user"', "Standard platform user. PRISM coaching, chat, documents, goals.", "MAIN, TOOLS, ACCOUNT"],
        ["Manager", '"manager"', "Team leader. Manages team members, hiring, training, career development.", "MAIN, MANAGER, TOOLS, ACCOUNT"],
        ["Company Admin", '"company-admin"', "Organization administrator. Org chart, company-wide training, assignments.", "MAIN, COMPANY, TOOLS, ACCOUNT"],
        ["Practitioner", '"practitioner"', "PRISM coach/practitioner. Manages clients, coaching sessions, credits.", "MAIN, PRACTITIONER, TOOLS, ACCOUNT"],
        ["Distributor", '"distributor"', "Credit distributor. Manages practitioners, purchases and allocates credits.", "MAIN, DISTRIBUTOR, TOOLS, ACCOUNT"],
        ["Super Admin", '"super-admin"', "Platform administrator. (Phase 2 \u2014 will add ADMIN section later.)", "ALL sections + ADMIN"],
    ],
)

doc.add_heading("4.2 Multi-Role Logic", level=2)
doc.add_paragraph(
    "Users can hold multiple roles simultaneously. When a user logs in, the system checks "
    "all roles assigned to them and renders all corresponding sidebar navigation groups."
)
doc.add_paragraph("Examples:")
bullet("A user with roles [user, manager] sees: MAIN + MANAGER + TOOLS + ACCOUNT")
bullet("A user with roles [user, company-admin, manager] sees: MAIN + MANAGER + COMPANY + TOOLS + ACCOUNT")
bullet("A user with roles [practitioner, distributor] sees: MAIN + PRACTITIONER + DISTRIBUTOR + TOOLS + ACCOUNT")

doc.add_heading("4.3 Implementation: Multi-Role in Auth", level=2)
doc.add_paragraph("Changes required to support multi-role:")
bullet("Backend API must return roles as an array: roles: [\"user\", \"manager\"] instead of role: \"user\"")
bullet("AuthContext: Store roles as string[] instead of string. Update completeAuthFromPayload().")
bullet("STORAGE_KEYS: Store USER_ROLES (array) instead of USER_ROLE (string)")
bullet("ProtectedRoute: Update hasAccess() to check if ANY of the user\u2019s roles grants access")
bullet("Sidebar: Iterate over user\u2019s roles, collect all applicable nav groups, render unified sidebar")

doc.add_heading("4.4 Default Dashboard Selection", level=2)
doc.add_paragraph(
    "When a user has multiple roles, the landing page should show the dashboard for their "
    "primary (highest-privilege) role. Role priority order: super-admin > distributor > "
    "practitioner > company-admin > manager > user. The user can navigate to other sections freely."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. USER DASHBOARD
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Section 1: User Dashboard \u2014 Detailed Breakdown", level=1)
doc.add_paragraph("The MAIN section is visible to ALL users regardless of role.")

doc.add_heading("5.1 Dashboard (/dashboard)", level=2)
doc.add_paragraph("Primary user home. From the new UI prototype:")
bullet("Welcome banner with user\u2019s name, motivational text")
bullet("Quick stat cards: Active Goals, Sessions This Month, PRISM Score, Documents")
bullet("Recent activity feed")
bullet("Quick-action buttons: Start Coaching Session, View PRISM Report, Upload Document")
bullet("PRISM profile summary card with behavioral color distribution (Red/Blue/Green/Yellow)")

doc.add_heading("5.2 Chat (/chat)", level=2)
doc.add_paragraph("Existing coach chat functionality, adapted to new layout:")
bullet("Coach selector sidebar (existing conversation list)")
bullet("Full chat window with text + voice (existing ChatWindow component)")
bullet("Message feedback widget (existing MessageFeedback component)")
bullet("Document sharing panel (existing DocumentsSidePanel)")

doc.add_heading("5.3 PRISM Report (/prism-report)", level=2)
doc.add_paragraph("NEW page \u2014 individual PRISM behavioral assessment report:")
bullet("PRISM color wheel visualization (Red=Drive, Blue=Compliance, Green=Steadiness, Yellow=Influence)")
bullet("Behavioral scores breakdown with percentages")
bullet("Strengths and development areas")
bullet("Communication style analysis")
bullet("Team compatibility insights")
bullet("Historical score trends")

doc.add_heading("5.4 Documents (/documents)", level=2)
doc.add_paragraph("Existing documents page. Minimal changes needed:")
bullet("File upload, download, delete, preview (existing)")
bullet("File type icons (PDF, DOC, XLS, IMG, FILE) \u2014 add type-specific icons per new UI")

doc.add_heading("5.5 Goals (/goals)", level=2)
doc.add_paragraph("NEW page \u2014 personal goal tracking:")
bullet("Set New Goal button \u2192 modal with goal form")
bullet("Goal cards with progress bars, due dates, status (Active, Completed, Overdue)")
bullet("Filter by status: All, Active, Completed, Overdue")
bullet("Goal categories tied to PRISM development areas")
bullet("Metrics: Active Goals count, Goals Completed count, Completion Rate")

doc.add_heading("5.6 Mapping to Existing Functionality", level=2)
add_table(
    ["New UI Element", "Existing Component/Page", "Action"],
    [
        ["Dashboard", "Home.tsx", "REDESIGN \u2014 replace welcome banner with new stat cards + PRISM summary"],
        ["Chat", "Dashboard.tsx + CoachChat.tsx", "REUSE \u2014 move into /chat route, same ChatWindow component"],
        ["PRISM Report", "None", "NEW BUILD \u2014 entire page + PRISM visualization components"],
        ["Documents", "Documents.tsx", "MINOR MODIFY \u2014 add file type icons, match new card styling"],
        ["Goals", "None", "NEW BUILD \u2014 entire page + goal service/hook/types"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. MANAGER DASHBOARD
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Section 2: Manager Dashboard \u2014 Detailed Breakdown", level=1)

doc.add_heading("6.1 Manager Dashboard (/manager/dashboard)", level=2)
doc.add_paragraph("Team Manager overview page:")
bullet("KPI cards: Team Members, Active Teams, PRISM Assessed %, Training Completion %")
bullet("Team Health Summary chart (engagement metrics)")
bullet("My Team\u2019s Progress chart")
bullet("Quick actions: Schedule Team Assessment, Start Coaching Session, View Reports")

doc.add_heading("6.2 My Team (/manager/team)", level=2)
bullet("Team member list table: Name, Role, PRISM Status, Sessions, Last Active")
bullet("Team Assessment overview with PRISM color distribution per member")
bullet("Team Balance visualization (PRISM quadrant balance)")
bullet("Team Compatibility Insights")
bullet("Team Dynamics & PRISM analysis")
bullet("Click member \u2192 drill down to individual profile with PRISM data")

doc.add_heading("6.3 Career Management (/manager/career)", level=2)
bullet("Career development plans per team member")
bullet("Career Plans list with status tracking")
bullet("Skills assessment matrix")
bullet("Development goals linked to PRISM behavioral insights")

doc.add_heading("6.4 Hiring (/manager/hiring)", level=2)
bullet("Hiring Pipeline board view (Kanban-style or table)")
bullet("KPIs: Open Positions, Total Candidates, Interviews This Week, Avg Time to Hire")
bullet("Position cards with candidate count and pipeline stage")

doc.add_heading("6.5 Team Building (/manager/team-building)", level=2)
bullet("PRISM Team Workshop planner")
bullet("Team Building Exercise scheduler")
bullet("Team dynamics analysis tools")
bullet("Workshop templates based on PRISM profiles")

doc.add_heading("6.6 Training (/manager/training)", level=2)
bullet("Training Programs list: Leadership Essentials, Technical Excellence, Compliance, etc.")
bullet("Program cards: enrolled count, completion rate, status")
bullet("Assign training to team members")
bullet("Track completion per individual")

doc.add_heading("6.7 Leadership (/manager/leadership)", level=2)
bullet("Leadership development profiles for team members")
bullet("Leadership assessment scores")
bullet("Leadership training recommendations based on PRISM")

doc.add_heading("6.8 Interviews (/manager/interviews)", level=2)
bullet("Interview schedule (calendar or list view)")
bullet("Candidate profiles with PRISM assessment results")
bullet("Interview feedback forms")
bullet("Metrics: Interviews This Week, Completion Rate")

doc.add_heading("6.9 Candidates (/manager/candidates)", level=2)
bullet("Candidate list with filtering (status: New, Assessment, Interview, Review)")
bullet("Candidate profile cards with PRISM assessment preview")
bullet("Candidate pipeline tracking")

doc.add_heading("6.10 Job DNA (/manager/job-dna)", level=2)
bullet("Job DNA profile builder \u2014 define ideal PRISM behavioral profile for a role")
bullet("PRISM-based job matching score")
bullet("Compare candidate PRISM profiles against Job DNA template")

doc.add_heading("6.11 Mapping to Existing Functionality", level=2)
add_table(
    ["New UI Element", "Existing", "Action"],
    [
        ["Manager Dashboard", "SA Dashboard.tsx (partial)", "NEW BUILD \u2014 manager-specific KPIs, team health"],
        ["My Team", "SA UserManagement.tsx (partial)", "MAJOR NEW BUILD \u2014 PRISM team analytics, member profiles"],
        ["Career Mgmt", "None", "NEW BUILD"],
        ["Hiring", "None", "NEW BUILD \u2014 hiring pipeline, candidate tracking"],
        ["Team Building", "None", "NEW BUILD \u2014 PRISM workshop tools"],
        ["Training", "None", "NEW BUILD \u2014 training program management"],
        ["Leadership", "None", "NEW BUILD"],
        ["Interviews", "None", "NEW BUILD"],
        ["Candidates", "None", "NEW BUILD"],
        ["Job DNA", "None", "NEW BUILD \u2014 PRISM-based job profiling"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. COMPANY ADMIN DASHBOARD
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Section 3: Company Admin Dashboard \u2014 Detailed Breakdown", level=1)

doc.add_heading("7.1 Company Dashboard (/company/dashboard)", level=2)
bullet("KPIs: Total Employees, Active Teams, PRISM Assessed %, Training Completion %")
bullet("Company-wide engagement metrics")
bullet("Department comparison charts")
bullet("PRISM distribution across the entire organization")

doc.add_heading("7.2 Org Chart (/company/org-chart)", level=2)
bullet("Interactive organizational hierarchy visualization")
bullet("Department tree: CEO \u2192 VPs \u2192 Directors \u2192 Managers \u2192 ICs")
bullet("Click any node to view profile, PRISM data, team info")
bullet("Roles from prototype: CEO, VP Engineering, VP Marketing, VP HR, VP Operations, Directors, Leads")

doc.add_heading("7.3 Teams (/company/teams)", level=2)
bullet("Team directory with search and filter")
bullet("Team cards: name, member count, PRISM balance, engagement rate")
bullet("Team Management: create, edit, assign members")
bullet("Team Assessments overview across all teams")

doc.add_heading("7.4 Training (/company/training)", level=2)
bullet("Company-wide training program management")
bullet("Programs: Leadership Accelerator, Communication Skills, Compliance & Ethics, etc.")
bullet("Enrollment tracking, completion rates, program effectiveness metrics")
bullet("Assign programs to departments/teams")

doc.add_heading("7.5 Assignments (/company/assignments)", level=2)
bullet("Coach-to-team/department assignment management")
bullet("PRISM assessment assignment scheduling (quarterly cycles)")
bullet("Training assignment tracking")
bullet("Bulk assignment operations")

doc.add_heading("7.6 Leadership Profiles (/company/leadership)", level=2)
bullet("Company-wide leadership profile database")
bullet("Leadership PRISM patterns across departments")
bullet("Leadership development tracking")

doc.add_heading("7.7 Mapping to Existing Functionality", level=2)
add_table(
    ["New UI Element", "Existing", "Action"],
    [
        ["Company Dashboard", "SA Dashboard.tsx", "REDESIGN \u2014 add PRISM org-wide metrics, department charts"],
        ["Org Chart", "SA OrganizationView.tsx (basic)", "MAJOR NEW BUILD \u2014 interactive tree visualization"],
        ["Teams", "SA TeamManagement.tsx (commented out)", "REACTIVATE + MAJOR ENHANCE \u2014 add PRISM team analytics"],
        ["Training", "None", "NEW BUILD"],
        ["Assignments", "SA OrganizationManagement.tsx (coach assign)", "ENHANCE \u2014 expand to PRISM + training assignments"],
        ["Leadership", "None", "NEW BUILD"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. PRACTITIONER DASHBOARD
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Section 4: Practitioner Dashboard \u2014 Detailed Breakdown", level=1)

doc.add_heading("8.1 Practitioner Home (/practitioner/dashboard)", level=2)
bullet("KPIs: Active Clients, Sessions This Month, Credits Remaining, Avg Client Satisfaction")
bullet("Upcoming sessions calendar widget")
bullet("Recent client activity feed")
bullet("Follow-up reminders (overdue items highlighted)")

doc.add_heading("8.2 My Clients (/practitioner/clients)", level=2)
bullet("Client list with search, filter by status (Active, Inactive, Pending Assessment)")
bullet("Client cards: name, org, PRISM status, last session, satisfaction score")
bullet("Click client \u2192 detailed profile with PRISM results, session history, notes")
bullet("Add new client workflow")

doc.add_heading("8.3 Conversations (/practitioner/conversations)", level=2)
bullet("Conversation history across all clients")
bullet("Filter by client, date range, coach agent used")
bullet("View conversation transcripts")
bullet("Practitioner Notes \u2014 add notes to conversations for follow-up")

doc.add_heading("8.4 Follow-Up (/practitioner/follow-up)", level=2)
bullet("Follow-up task list with due dates")
bullet("Status: Due Today, Overdue, Upcoming, Completed")
bullet("Create follow-up reminders linked to client sessions")
bullet("Email/notification integration for follow-up reminders")

doc.add_heading("8.5 Credits (/practitioner/credits)", level=2)
bullet("Credit balance display: Credits Remaining, Credits Used, Total Allocated")
bullet("Credit usage history table")
bullet("Credit consumption per session/client breakdown")
bullet("Request more credits from distributor")

doc.add_heading("8.6 Mapping to Existing Functionality", level=2)
add_table(
    ["New UI Element", "Existing", "Action"],
    [
        ["Practitioner Home", "None", "NEW BUILD \u2014 practitioner-specific dashboard"],
        ["My Clients", "None", "NEW BUILD \u2014 client management system"],
        ["Conversations", "Coach Chat (partial pattern)", "NEW BUILD \u2014 reuse ChatWindow components for viewing transcripts"],
        ["Follow-Up", "None", "NEW BUILD"],
        ["Credits", "None", "NEW BUILD \u2014 credit tracking system"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. DISTRIBUTOR DASHBOARD
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Section 5: Distributor Dashboard \u2014 Detailed Breakdown", level=1)

doc.add_heading("9.1 Distributor Home (/distributor/dashboard)", level=2)
bullet("KPIs: Total Practitioners, Active Practitioners, Total Credits, Credits Allocated")
bullet("Credit distribution overview chart")
bullet("Practitioner performance summary")
bullet("Revenue/usage trends")

doc.add_heading("9.2 Practitioners (/distributor/practitioners)", level=2)
bullet("Practitioner roster: name, status (Active/Inactive), credits allocated, clients served")
bullet("Add Practitioner workflow \u2192 modal or wizard")
bullet("Deactivate practitioner (credits returned to pool)")
bullet("View practitioner details: performance metrics, client list, credit usage")
bullet("Practitioner Management actions: allocate credits, adjust status")

doc.add_heading("9.3 Credits (/distributor/credits)", level=2)
bullet("Credit pool management: Total Purchased, Total Allocated, Available Pool")
bullet("Purchase Credits workflow")
bullet("Credit transaction history table")
bullet("Credit expiration tracking")
bullet("Usage analytics: credits used by practitioners, by time period")

doc.add_heading("9.4 Allocate Credits (/distributor/credit-allocation)", level=2)
bullet("Allocation interface: select practitioner(s), enter credit amount, confirm")
bullet("Bulk allocation to multiple practitioners")
bullet("Allocation history with status (allocated, used, returned)")
bullet("Ready to allocate summary card")

doc.add_heading("9.5 Mapping to Existing Functionality", level=2)
add_table(
    ["New UI Element", "Existing", "Action"],
    [
        ["Distributor Home", "None", "NEW BUILD"],
        ["Practitioners", "SA UserManagement.tsx (pattern only)", "NEW BUILD \u2014 reuse DataTable pattern"],
        ["Credits", "SA LicenceDetailsPage.tsx (partial)", "NEW BUILD \u2014 credit-based licensing system"],
        ["Allocate Credits", "None", "NEW BUILD"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. SUPER ADMIN (FUTURE)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Section 6: Super Admin (Future)", level=1)
doc.add_paragraph(
    "The Super Admin section will be added as a 6th section in a later phase. It will "
    "consolidate the existing SuperAdminLayout functionality into the unified sidebar under "
    "an ADMIN navigation group."
)
doc.add_paragraph("Planned items to carry forward from existing SuperAdminLayout:")
bullet("Dashboard (system stats, org/user overview)")
bullet("User Management (invite, edit, delete, role assignment)")
bullet("Coach Management (create, edit, deactivate)")
bullet("RLHF Training (feedback metrics, rating charts)")
bullet("Prompt Builder (system prompt wizard)")
bullet("Audit Log (event logging viewer)")
bullet("Organization Management (org CRUD, coach assignment)")
bullet("Settings")
bullet("Project Log")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. EXISTING FUNCTIONALITY MAPPING
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Existing Functionality Mapping", level=1)
doc.add_paragraph(
    "Complete mapping of every existing page/feature to the new UI, with disposition."
)

add_table(
    ["Current Feature", "Current Route", "New Route", "Disposition"],
    [
        ["Home (welcome banner)", "/home", "/dashboard", "REDESIGN \u2014 new stat cards, PRISM summary, activity feed"],
        ["Coach Chat List", "/dashboard", "/chat", "MOVE \u2014 relocate to /chat, same component"],
        ["Coach Chat Window", "/dashboard/:coach/chat", "/chat/:coach", "MOVE \u2014 same ChatWindow, new route"],
        ["Manage Coaches", "/coaches", "/settings/coaches", "MOVE \u2014 integrate into settings or keep as sub-page"],
        ["Documents", "/documents", "/documents", "KEEP \u2014 minor styling updates"],
        ["Settings", "/settings", "/settings", "KEEP \u2014 minor styling updates"],
        ["Help & Support", "/help", "Removed or /settings/help", "EVALUATE \u2014 may merge into settings"],
        ["Onboarding Flow", "/onboarding/*", "/onboarding/*", "KEEP \u2014 no change needed"],
        ["Auth Pages", "/login, /signup, etc.", "Same", "KEEP \u2014 no change needed"],
        ["SA Dashboard", "/super-admin/dashboard", "/admin/dashboard (Phase 2)", "DEFER \u2014 Super Admin section later"],
        ["SA User Mgmt", "/super-admin/users", "/admin/users (Phase 2)", "DEFER"],
        ["SA Coach Mgmt", "/super-admin/coaches", "/admin/coaches (Phase 2)", "DEFER"],
        ["SA RLHF Training", "/super-admin/rlhf-training", "/admin/rlhf (Phase 2)", "DEFER"],
        ["SA Prompt Builder", "/super-admin/prompt-builder", "/admin/prompts (Phase 2)", "DEFER"],
        ["SA Audit Log", "/super-admin/audit-log", "/admin/audit (Phase 2)", "DEFER"],
        ["SA Organizations", "/super-admin/organizations", "/admin/orgs (Phase 2)", "DEFER"],
        ["SA Settings", "/super-admin/settings", "/admin/settings (Phase 2)", "DEFER"],
        ["MessageFeedback", "Component in chat", "Same component in /chat", "KEEP"],
        ["Alex Floating", "Component in UserLayout", "Component in UnifiedLayout", "MOVE"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. NEW FEATURES REQUIRED
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("12. New Features Required", level=1)

doc.add_paragraph("Features that do NOT exist in the current codebase and must be built from scratch:")

add_table(
    ["Feature", "Section", "Complexity", "Backend API Required"],
    [
        ["PRISM Report Page", "User", "High", "Yes \u2014 PRISM assessment results API"],
        ["Goals System", "User", "Medium", "Yes \u2014 CRUD goals API"],
        ["PRISM Visualization Components", "User/Manager/Company", "High", "Yes \u2014 PRISM scoring data"],
        ["Manager Dashboard", "Manager", "Medium", "Yes \u2014 team analytics API"],
        ["Team Member PRISM Profiles", "Manager", "High", "Yes \u2014 team PRISM data"],
        ["Career Management", "Manager", "Medium", "Yes \u2014 career plans API"],
        ["Hiring Pipeline", "Manager", "High", "Yes \u2014 candidates/positions API"],
        ["Job DNA Profiler", "Manager", "High", "Yes \u2014 job profile PRISM matching API"],
        ["Interview Scheduling", "Manager", "Medium", "Yes \u2014 interview API"],
        ["Training Programs", "Manager/Company", "Medium", "Yes \u2014 training CRUD API"],
        ["Team Building Workshop Tools", "Manager", "Medium", "Yes \u2014 workshop API"],
        ["Org Chart Visualization", "Company", "High", "Yes \u2014 org hierarchy API"],
        ["Company Training Dashboard", "Company", "Medium", "Reuse training API"],
        ["Assignment Management", "Company", "Medium", "Enhance existing org API"],
        ["Practitioner Client Management", "Practitioner", "High", "Yes \u2014 practitioner-client API"],
        ["Practitioner Follow-Up System", "Practitioner", "Medium", "Yes \u2014 follow-up tasks API"],
        ["Credit System", "Practitioner/Distributor", "High", "Yes \u2014 credit ledger API"],
        ["Distributor Practitioner Mgmt", "Distributor", "Medium", "Yes \u2014 distributor API"],
        ["Credit Purchase Workflow", "Distributor", "High", "Yes \u2014 payment/credit purchase API"],
        ["Credit Allocation System", "Distributor", "Medium", "Yes \u2014 credit allocation API"],
        ["Notifications Page", "All", "Medium", "Yes \u2014 notifications API"],
        ["Messages Page", "All", "Medium", "Yes \u2014 messaging API"],
        ["Cost Dashboard", "Admin/Company", "High", "Yes \u2014 cost analytics API"],
        ["Multi-Role Auth", "All", "Medium", "Yes \u2014 modify auth response to return roles[]"],
        ["Unified Sidebar Layout", "All", "Medium", "No \u2014 frontend only"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 13. COMPONENTS TO MODIFY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Components to Modify", level=1)

add_table(
    ["File", "Modification", "Effort"],
    [
        ["src/context/AuthContext.tsx", "Store roles as string[] instead of string. Update completeAuthFromPayload, logout, role checks.", "Medium"],
        ["src/types/roles.ts", "Add new roles (manager, practitioner, distributor). Update ROLE_PERMISSIONS for multi-role. Update hasAccess() to check any role match.", "Medium"],
        ["src/constants/routes.ts", "Add all new route constants for manager/*, company/*, practitioner/*, distributor/* sections. Expand ROLES enum.", "Low"],
        ["src/routes.tsx", "Add ~30 new route entries for all 5 sections. Replace UserLayout/SuperAdminLayout with UnifiedLayout.", "Medium"],
        ["src/components/ProtectedRoute.tsx", "Update to handle roles[] array, check multi-role access.", "Low"],
        ["src/lib/secureStorage.ts", "Store USER_ROLES as JSON array instead of string.", "Low"],
        ["src/pages/user/Home.tsx", "REDESIGN as /dashboard with new stat cards, PRISM summary, activity feed.", "High"],
        ["src/pages/user/Dashboard.tsx", "MOVE to /chat route. Rename component.", "Low"],
        ["src/pages/user/Documents.tsx", "Minor styling updates: file type icons, card design match.", "Low"],
        ["src/pages/user/Settings.tsx", "Minor styling updates to match new design.", "Low"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 14. NEW COMPONENTS TO BUILD
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("14. New Components to Build", level=1)

doc.add_heading("14.1 Layout Components", level=2)
add_table(
    ["Component", "Path", "Description"],
    [
        ["UnifiedLayout", "src/layouts/UnifiedLayout.tsx", "Single layout replacing UserLayout + SuperAdminLayout. Renders role-based sidebar groups."],
        ["UnifiedSidebar", "src/components/layout/UnifiedSidebar.tsx", "Role-aware sidebar with grouped navigation. Receives user roles[], renders applicable groups."],
        ["SidebarNavGroup", "src/components/layout/SidebarNavGroup.tsx", "Collapsible navigation group (title + items). Reuse for MAIN, MANAGER, COMPANY, etc."],
        ["BreadcrumbHeader", "src/components/layout/BreadcrumbHeader.tsx", "Header with breadcrumb trail, date filter, export button."],
    ],
)

doc.add_heading("14.2 PRISM Components", level=2)
add_table(
    ["Component", "Path", "Description"],
    [
        ["PrismColorWheel", "src/components/prism/PrismColorWheel.tsx", "PRISM quadrant visualization (Red/Blue/Green/Yellow). SVG or Canvas based."],
        ["PrismScoreCard", "src/components/prism/PrismScoreCard.tsx", "Individual PRISM dimension score with bar chart."],
        ["PrismTeamBalance", "src/components/prism/PrismTeamBalance.tsx", "Team PRISM distribution chart showing balance across dimensions."],
        ["PrismCompatibility", "src/components/prism/PrismCompatibility.tsx", "Two-person compatibility analysis based on PRISM profiles."],
        ["PrismAssessment", "src/components/prism/PrismAssessment.tsx", "Assessment form/wizard for taking PRISM assessment."],
    ],
)

doc.add_heading("14.3 Manager Components", level=2)
add_table(
    ["Component", "Path", "Description"],
    [
        ["TeamMemberTable", "src/components/manager/TeamMemberTable.tsx", "DataTable of team members with PRISM status, sessions, last active."],
        ["HiringPipeline", "src/components/manager/HiringPipeline.tsx", "Kanban or table view of hiring pipeline stages."],
        ["CandidateCard", "src/components/manager/CandidateCard.tsx", "Candidate profile card with PRISM preview."],
        ["JobDnaBuilder", "src/components/manager/JobDnaBuilder.tsx", "Interactive PRISM-based job profile builder."],
        ["TrainingProgramCard", "src/components/manager/TrainingProgramCard.tsx", "Training program card with enrollment and completion metrics."],
        ["InterviewSchedule", "src/components/manager/InterviewSchedule.tsx", "Calendar/list view of scheduled interviews."],
    ],
)

doc.add_heading("14.4 Company Admin Components", level=2)
add_table(
    ["Component", "Path", "Description"],
    [
        ["OrgChartTree", "src/components/company/OrgChartTree.tsx", "Interactive hierarchical org chart visualization."],
        ["DepartmentCard", "src/components/company/DepartmentCard.tsx", "Department summary card with member count, PRISM data."],
        ["CompanyTrainingDashboard", "src/components/company/CompanyTrainingDashboard.tsx", "Company-wide training metrics and program management."],
    ],
)

doc.add_heading("14.5 Practitioner Components", level=2)
add_table(
    ["Component", "Path", "Description"],
    [
        ["ClientCard", "src/components/practitioner/ClientCard.tsx", "Client profile card with PRISM status, satisfaction score."],
        ["ClientDetailPanel", "src/components/practitioner/ClientDetailPanel.tsx", "Full client detail view: PRISM results, session history, notes."],
        ["FollowUpList", "src/components/practitioner/FollowUpList.tsx", "Follow-up task list with status filtering."],
        ["CreditBalance", "src/components/practitioner/CreditBalance.tsx", "Credit balance display with usage chart."],
        ["PractitionerNotes", "src/components/practitioner/PractitionerNotes.tsx", "Session notes editor tied to conversations."],
    ],
)

doc.add_heading("14.6 Distributor Components", level=2)
add_table(
    ["Component", "Path", "Description"],
    [
        ["PractitionerRoster", "src/components/distributor/PractitionerRoster.tsx", "Practitioner management table with credit allocation."],
        ["CreditPoolDashboard", "src/components/distributor/CreditPoolDashboard.tsx", "Credit pool overview with allocation chart."],
        ["CreditAllocationForm", "src/components/distributor/CreditAllocationForm.tsx", "Allocate credits to practitioners (form + confirmation)."],
        ["PurchaseCreditsModal", "src/components/distributor/PurchaseCreditsModal.tsx", "Purchase credits workflow."],
        ["AddPractitionerModal", "src/components/distributor/AddPractitionerModal.tsx", "Add new practitioner wizard."],
    ],
)

doc.add_heading("14.7 Shared Components", level=2)
add_table(
    ["Component", "Path", "Description"],
    [
        ["MetricCard", "src/components/shared/MetricCard.tsx", "Reusable KPI card (value, label, trend indicator). Used across all dashboards."],
        ["NotificationsPage", "src/pages/shared/Notifications.tsx", "Notification center with read/unread, filters."],
        ["MessagesPage", "src/pages/shared/Messages.tsx", "Internal messaging between users/practitioners/managers."],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 15. ROUTING & NAVIGATION CHANGES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("15. Routing & Navigation Changes", level=1)

doc.add_heading("15.1 New Route Definitions", level=2)
add_table(
    ["Section", "Route", "Page Component", "Status"],
    [
        ["User", "/dashboard", "UserDashboard.tsx", "Redesign Home.tsx"],
        ["User", "/chat", "ChatHub.tsx", "Move from /dashboard"],
        ["User", "/chat/:coachId", "CoachChat.tsx", "Move from /dashboard/:coach/chat"],
        ["User", "/prism-report", "PrismReport.tsx", "NEW"],
        ["User", "/documents", "Documents.tsx", "Keep"],
        ["User", "/goals", "Goals.tsx", "NEW"],
        ["Manager", "/manager/dashboard", "ManagerDashboard.tsx", "NEW"],
        ["Manager", "/manager/team", "ManagerTeam.tsx", "NEW"],
        ["Manager", "/manager/career", "ManagerCareer.tsx", "NEW"],
        ["Manager", "/manager/hiring", "ManagerHiring.tsx", "NEW"],
        ["Manager", "/manager/team-building", "ManagerTeamBuilding.tsx", "NEW"],
        ["Manager", "/manager/training", "ManagerTraining.tsx", "NEW"],
        ["Manager", "/manager/leadership", "ManagerLeadership.tsx", "NEW"],
        ["Manager", "/manager/interviews", "ManagerInterviews.tsx", "NEW"],
        ["Manager", "/manager/candidates", "ManagerCandidates.tsx", "NEW"],
        ["Manager", "/manager/job-dna", "ManagerJobDna.tsx", "NEW"],
        ["Company", "/company/dashboard", "CompanyDashboard.tsx", "NEW"],
        ["Company", "/company/org-chart", "CompanyOrgChart.tsx", "NEW"],
        ["Company", "/company/teams", "CompanyTeams.tsx", "NEW"],
        ["Company", "/company/training", "CompanyTraining.tsx", "NEW"],
        ["Company", "/company/assignments", "CompanyAssignments.tsx", "NEW"],
        ["Company", "/company/leadership", "CompanyLeadership.tsx", "NEW"],
        ["Practitioner", "/practitioner/dashboard", "PractitionerDashboard.tsx", "NEW"],
        ["Practitioner", "/practitioner/clients", "PractitionerClients.tsx", "NEW"],
        ["Practitioner", "/practitioner/conversations", "PractitionerConversations.tsx", "NEW"],
        ["Practitioner", "/practitioner/follow-up", "PractitionerFollowUp.tsx", "NEW"],
        ["Practitioner", "/practitioner/credits", "PractitionerCredits.tsx", "NEW"],
        ["Distributor", "/distributor/dashboard", "DistributorDashboard.tsx", "NEW"],
        ["Distributor", "/distributor/practitioners", "DistributorPractitioners.tsx", "NEW"],
        ["Distributor", "/distributor/credits", "DistributorCredits.tsx", "NEW"],
        ["Distributor", "/distributor/credit-allocation", "DistributorAllocation.tsx", "NEW"],
        ["Shared", "/settings", "Settings.tsx", "Keep"],
        ["Shared", "/notifications", "Notifications.tsx", "NEW"],
        ["Shared", "/messages", "Messages.tsx", "NEW"],
        ["Shared", "/assessment", "Assessment.tsx", "NEW"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 16. DATA LAYER
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("16. Data Layer: New Services, Hooks & Types", level=1)

doc.add_paragraph("New backend API domains required to support the new UI:")

add_table(
    ["Domain", "Service File", "Hook File", "Key Endpoints"],
    [
        ["PRISM", "src/services/prism/prism.service.ts", "src/hooks/prism/usePrism.ts", "GET /v1/prism/assessment, POST /v1/prism/assessment, GET /v1/prism/report/{userId}"],
        ["Goals", "src/services/goals/goals.service.ts", "src/hooks/goals/useGoals.ts", "CRUD /v1/goals, GET /v1/goals/stats"],
        ["Team", "src/services/team/team.service.ts", "src/hooks/team/useTeam.ts", "GET /v1/team/members, GET /v1/team/analytics, GET /v1/team/prism-balance"],
        ["Career", "src/services/career/career.service.ts", "src/hooks/career/useCareer.ts", "CRUD /v1/career/plans"],
        ["Hiring", "src/services/hiring/hiring.service.ts", "src/hooks/hiring/useHiring.ts", "CRUD /v1/hiring/positions, /v1/hiring/candidates"],
        ["Training", "src/services/training/training.service.ts", "src/hooks/training/useTraining.ts", "CRUD /v1/training/programs, /v1/training/enrollment"],
        ["Interviews", "src/services/interviews/interviews.service.ts", "src/hooks/interviews/useInterviews.ts", "CRUD /v1/interviews"],
        ["Job DNA", "src/services/job-dna/job-dna.service.ts", "src/hooks/job-dna/useJobDna.ts", "CRUD /v1/job-dna/profiles, GET /v1/job-dna/match"],
        ["Org Chart", "src/services/org-chart/org-chart.service.ts", "src/hooks/org-chart/useOrgChart.ts", "GET /v1/org-chart/hierarchy"],
        ["Practitioner", "src/services/practitioner/practitioner.service.ts", "src/hooks/practitioner/usePractitioner.ts", "CRUD /v1/practitioner/clients, /v1/practitioner/follow-ups"],
        ["Credits", "src/services/credits/credits.service.ts", "src/hooks/credits/useCredits.ts", "GET /v1/credits/balance, POST /v1/credits/allocate, POST /v1/credits/purchase"],
        ["Distributor", "src/services/distributor/distributor.service.ts", "src/hooks/distributor/useDistributor.ts", "GET /v1/distributor/practitioners, CRUD operations"],
        ["Notifications", "src/services/notifications/notifications.service.ts", "src/hooks/notifications/useNotifications.ts", "GET /v1/notifications, PATCH /v1/notifications/{id}/read"],
        ["Messages", "src/services/messages/messages.service.ts", "src/hooks/messages/useMessages.ts", "GET /v1/messages, POST /v1/messages"],
        ["Cost Analytics", "src/services/cost/cost.service.ts", "src/hooks/cost/useCost.ts", "GET /v1/cost/platform, /v1/cost/company/{id}, /v1/cost/team/{id}, /v1/cost/user/{id}"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 17. COST DASHBOARD INTEGRATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("17. Cost Dashboard Integration", level=1)
doc.add_paragraph(
    "The cost dashboard (cost-dashboard.html) provides a drill-down cost analytics view "
    "with 4 hierarchical levels: Platform \u2192 Company \u2192 Team \u2192 User."
)

doc.add_heading("17.1 Cost Dashboard Levels", level=2)
add_table(
    ["Level", "View", "KPIs", "Available To"],
    [
        ["Platform", "Platform Cost Overview", "Total Cost MTD, Active Users, New Users, Cost/User, Total Sessions", "Super Admin, Distributor"],
        ["Company", "Company Cost Analysis", "Org Cost, Active Users, Cost/User, Sessions, Department breakdown table", "Company Admin"],
        ["Team", "Team Cost Analysis", "Team Cost, Team Members, Engagement Rate, Avg Sessions/User, Member table", "Manager"],
        ["User", "User Cost Profile", "Sessions, Total Tokens, Audio Minutes, Cost/Session, Cost breakdown bars", "User (own data), Manager (team members)"],
    ],
)

doc.add_heading("17.2 Cost Dashboard Components", level=2)
add_table(
    ["Component", "Path", "Description"],
    [
        ["CostPlatformDashboard", "src/components/cost/CostPlatformDashboard.tsx", "Platform-level KPIs, Cost by Week pie, Cost Distribution pie, Agent Usage pie, Top Orgs table, Top Users table"],
        ["CostCompanyDashboard", "src/components/cost/CostCompanyDashboard.tsx", "Company KPIs, Department Cost Breakdown table (sortable)"],
        ["CostTeamDashboard", "src/components/cost/CostTeamDashboard.tsx", "Team KPIs, Team Members table with avatar, role, sessions, tokens, cost, last active"],
        ["CostUserProfile", "src/components/cost/CostUserProfile.tsx", "User profile card, metric cards, cost breakdown progress bars (Tokens, Audio, PRISM, Embeddings)"],
        ["CostBreadcrumb", "src/components/cost/CostBreadcrumb.tsx", "Drill-down breadcrumb: Platform > Company > Team > User"],
    ],
)

doc.add_heading("17.3 Cost Dashboard Charts", level=2)
bullet("Cost by Week \u2014 Pie chart (Recharts PieChart, 4 weeks)")
bullet("Cost Distribution \u2014 Pie chart (AI 70%, Infrastructure 18%, 3rd Party 12%)")
bullet("Cost by Agent \u2014 Pie chart (Meridian, Aura, Nova, Atlas, Echo, Other)")
bullet("All charts use the teal (#00A3A3) + green + blue + orange color palette")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 18. IMPLEMENTATION PHASES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("18. Implementation Phases", level=1)

add_table(
    ["Phase", "Duration", "Deliverables", "Effort"],
    [
        ["Phase 1: Unified Layout + Multi-Role Auth", "2\u20133 weeks",
         "UnifiedLayout, UnifiedSidebar, BreadcrumbHeader, multi-role AuthContext, updated ProtectedRoute, role-based nav group rendering, route restructuring",
         "Medium \u2014 foundational, must be done first"],
        ["Phase 2: User Section (MAIN)", "2\u20133 weeks",
         "Redesigned Dashboard, Chat move to /chat, PRISM Report page, Goals page, PRISM visualization components",
         "High \u2014 PRISM components are complex"],
        ["Phase 3: Manager Section", "3\u20134 weeks",
         "Manager Dashboard, My Team, Career Mgmt, Hiring Pipeline, Team Building, Training, Leadership, Interviews, Candidates, Job DNA",
         "High \u2014 10 new pages, heavy backend API dependency"],
        ["Phase 4: Company Admin Section", "2\u20133 weeks",
         "Company Dashboard, Org Chart, Teams, Training, Assignments, Leadership Profiles",
         "High \u2014 Org Chart visualization is complex"],
        ["Phase 5: Practitioner Section", "2\u20133 weeks",
         "Practitioner Dashboard, My Clients, Conversations, Follow-Up, Credits",
         "Medium-High \u2014 client management system"],
        ["Phase 6: Distributor Section", "2 weeks",
         "Distributor Dashboard, Practitioners, Credits, Credit Allocation",
         "Medium \u2014 credit system + practitioner management"],
        ["Phase 7: Shared Features", "1\u20132 weeks",
         "Notifications page, Messages page, Assessment page, Cost Dashboard",
         "Medium"],
        ["Phase 8: Super Admin Section", "1\u20132 weeks",
         "Migrate existing SA pages into unified layout under ADMIN nav group",
         "Low \u2014 mostly moving existing code"],
        ["Phase 9: PWA + Polish", "1\u20132 weeks",
         "Service worker, manifest, offline banner, install prompt, final styling pass",
         "Medium"],
        ["TOTAL", "16\u201325 weeks", "", ""],
    ],
)

doc.add_paragraph()
bold_para("Solo developer estimate: 5\u20137 months", 12, RGBColor(0, 51, 102))
bold_para("Team of 2\u20133: 3\u20134 months", 12, RGBColor(0, 51, 102))

doc.add_heading("18.1 Critical Path", level=2)
doc.add_paragraph("Phase 1 (Layout + Auth) must be completed first \u2014 all other phases depend on it.")
doc.add_paragraph("Phases 2\u20136 can be done in parallel by different developers if the backend APIs are ready.")
doc.add_paragraph("Phase 8 (Super Admin migration) can start any time after Phase 1.")

doc.add_heading("18.2 Backend API Dependency", level=2)
doc.add_paragraph(
    "The new UI requires approximately 15 new API domains. Backend API development must run "
    "in parallel with or ahead of frontend work. The biggest blockers are:"
)
bullet("PRISM assessment/reporting API \u2014 blocks Phase 2 and touches Phases 3, 4, 5")
bullet("Credit system API \u2014 blocks Phases 5 and 6")
bullet("Hiring/Candidate/Interview API \u2014 blocks Phase 3")
bullet("Cost analytics API \u2014 blocks Phase 7")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 19. RISK ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("19. Risk Assessment", level=1)

add_table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Backend APIs not ready in time", "High", "Critical", "Build frontend with mock data. Use MSW (Mock Service Worker) for API mocking. Ship incrementally."],
        ["Multi-role auth introduces bugs", "Medium", "High", "Extensive testing of role combinations. Write unit tests for hasAccess() with all role combos."],
        ["Scope creep across 5 sections", "High", "High", "Ship one section at a time. Each phase is independently deployable."],
        ["PRISM visualization complexity", "Medium", "Medium", "Use existing charting libraries (Recharts). Consider d3.js only if Recharts insufficient."],
        ["Org Chart rendering at scale", "Medium", "Medium", "Use virtual rendering. Limit initial depth. Lazy-load subtrees on expand."],
        ["Credit system financial accuracy", "Medium", "Critical", "Server-side ledger with transactions. Frontend is read-only display. Backend enforces all credit logic."],
        ["Breaking existing user workflows", "Medium", "High", "Feature-flag the new UI. Let users opt-in initially. Run old and new layouts in parallel during transition."],
        ["Performance with 5+ dashboard sections", "Low", "Medium", "Code-split by route. Lazy-load section pages. Only load data for active section."],
        ["Mobile responsiveness", "Medium", "Medium", "New sidebar collapses to hamburger on mobile. Test across breakpoints."],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("\u2014 End of Document \u2014")
run.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── SAVE ────────────────────────────────────────────────────────────────
output_dir = "/Users/williambrown/Dropbox/AES Material/Inspire-X/New IG Projects/Local_IG-App_UI/inspire-genius-frontend"
output_path = os.path.join(output_dir, "Inspire_Genius_New_UI_Adoption_Plan.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
