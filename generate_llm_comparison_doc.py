"""Generate LLM Strategy Comparison Word document: Multi-Provider vs AWS Bedrock."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0, 51, 102)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def bold_paragraph(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Inspire Genius\nLLM Strategy Comparison")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    "Multi-Provider (OpenAI + Gemini + Voyage AI)\nvs.\nAWS Bedrock Standardization"
)
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0, 128, 128)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Prepared for: ").bold = True
meta.add_run("AES / Inspire Genius Team\n")
meta.add_run("Date: ").bold = True
meta.add_run("March 13, 2026\n")
meta.add_run("Classification: ").bold = True
meta.add_run("Internal \u2014 Strategic Planning")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Current State: Multi-Provider LLM Architecture",
    "3. Proposed State: AWS Bedrock Standardization",
    "4. Comparison: Cost",
    "5. Comparison: Availability & Reliability",
    "6. Comparison: Scalability",
    "7. Comparison: Flexibility",
    "8. Comparison: Maintainability",
    "9. Comparison: Extensibility",
    "10. Side-by-Side Summary Matrix",
    "11. Pros and Cons",
    "12. Benefits and Risks",
    "13. Migration Effort Estimate",
    "14. Opinion & Recommendation",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Number")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "Inspire Genius currently operates a multi-provider AI strategy across its backend. "
    "Google Gemini serves as the primary LLM for all AI coach conversations and reasoning. "
    "OpenAI provides speech-to-text (Whisper/gpt-4o-transcribe), text-to-speech (gpt-4o-mini-tts), "
    "and a secondary embeddings model (text-embedding-3-small). Google also provides embeddings "
    "(gemini-embedding-001) as the default for the Milvus vector database. Voyage AI has an API key "
    "configured but is not actively used. Anthropic/Claude models are not currently integrated in the backend."
)
doc.add_paragraph(
    "This document evaluates whether Inspire Genius should standardize on AWS Bedrock and AWS-native "
    "AI services (Amazon Transcribe for STT, Amazon Polly for TTS, Bedrock for LLM inference and embeddings) "
    "versus maintaining the current multi-provider approach. The comparison covers six dimensions: "
    "Cost, Availability, Scalability, Flexibility, Maintainability, and Extensibility."
)
doc.add_paragraph(
    "Bottom line: The current multi-provider strategy is the stronger choice for Inspire Genius today. "
    "AWS Bedrock offers meaningful advantages in operational simplicity and AWS ecosystem integration, "
    "but the trade-offs in model quality (especially TTS/STT), cost, and vendor lock-in outweigh those "
    "benefits at the platform\u2019s current scale and maturity. A hybrid approach\u2014keeping best-of-breed "
    "models for core AI features while using Bedrock for specific workloads\u2014is the recommended path forward."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. CURRENT STATE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Current State: Multi-Provider LLM Architecture", level=1)

doc.add_heading("2.1 Provider Map", level=2)
add_table(
    ["Function", "Provider", "Model", "Integration"],
    [
        ["Primary LLM (Reasoning)", "Google Gemini", "gemini-3-flash-preview", "google-genai SDK, direct API"],
        ["Query Refinement", "Google Gemini", "gemini-3-flash-preview", "Structured output via google-genai"],
        ["Audio Transcription (STT)", "OpenAI", "gpt-4o-transcribe (Whisper)", "openai SDK, async client"],
        ["Text-to-Speech (TTS) \u2014 Primary", "OpenAI", "gpt-4o-mini-tts", "openai SDK, streaming"],
        ["Text-to-Speech (TTS) \u2014 Secondary", "Google Gemini", "gemini-2.5-flash-preview-tts", "google-genai SDK"],
        ["Embeddings \u2014 Active", "Google", "gemini-embedding-001", "LangChain GoogleGenerativeAIEmbeddings"],
        ["Embeddings \u2014 Configured", "OpenAI", "text-embedding-3-small (512d)", "LangChain OpenAIEmbeddings"],
        ["Embeddings \u2014 Reserved", "Voyage AI", "(not active)", "API key configured, unused"],
        ["Vector Database", "Milvus", "milvus-lite", "LangChain Milvus wrapper, HNSW/L2"],
    ],
)

doc.add_heading("2.2 Architecture Flow", level=2)
doc.add_paragraph(
    "User speaks \u2192 WebSocket sends audio binary \u2192 OpenAI Whisper transcribes \u2192 "
    "Gemini generates response \u2192 OpenAI TTS or Gemini TTS streams audio back \u2192 "
    "WebSocket sends audio binary to client. For text-only chat, the STT/TTS steps are skipped."
)
doc.add_paragraph(
    "RAG (Retrieval-Augmented Generation) flow: Document uploaded \u2192 chunked \u2192 "
    "Google Embeddings generates vectors \u2192 stored in Milvus \u2192 at query time, "
    "query embedded \u2192 similar chunks retrieved \u2192 injected into Gemini prompt."
)

doc.add_heading("2.3 Dependencies", level=2)
add_table(
    ["Package", "Version", "Purpose"],
    [
        ["openai (async-openai)", ">=0.0.52", "OpenAI API client (STT, TTS, embeddings)"],
        ["google-generativeai", ">=0.8.6", "Google Gemini (legacy SDK)"],
        ["google-genai", ">=1.60.0", "Google Gemini (new SDK)"],
        ["langchain", ">=1.2.7", "LLM orchestration framework"],
        ["langchain-openai", ">=1.1.7", "OpenAI LangChain integration"],
        ["langchain-google-genai", ">=4.2.0", "Google LangChain integration"],
        ["langchain-milvus", ">=0.3.3", "Milvus vector DB integration"],
        ["milvus-lite", ">=2.5.1", "Embedded vector database"],
    ],
)

doc.add_heading("2.4 API Keys Required", level=2)
bullet("OPENAI_API_KEY \u2014 for STT, TTS, and backup embeddings")
bullet("GEMINI_API_KEY \u2014 for primary LLM reasoning and active embeddings")
bullet("VOYAGEAI_API_KEY \u2014 configured but unused")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. PROPOSED STATE: AWS BEDROCK
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Proposed State: AWS Bedrock Standardization", level=1)

doc.add_heading("3.1 AWS Service Mapping", level=2)
add_table(
    ["Function", "Current Provider", "AWS Replacement", "AWS Model/Service"],
    [
        ["Primary LLM", "Google Gemini", "Amazon Bedrock", "Claude 4 Sonnet (Anthropic), or Llama 4, or Amazon Nova"],
        ["Query Refinement", "Google Gemini", "Amazon Bedrock", "Claude 4 Haiku (fast, cheap)"],
        ["STT (Transcription)", "OpenAI Whisper", "Amazon Transcribe", "Real-time or batch transcription"],
        ["TTS (Speech)", "OpenAI / Gemini TTS", "Amazon Polly", "Neural voices, NTTS, Generative Engine"],
        ["Embeddings", "Google / OpenAI", "Amazon Bedrock", "Amazon Titan Embeddings v2, or Cohere Embed v3"],
        ["Vector Database", "Milvus (self-hosted)", "Amazon OpenSearch Serverless", "Vector search with k-NN plugin"],
        ["Orchestration", "LangChain", "Amazon Bedrock Agents / LangChain", "Bedrock-native or keep LangChain"],
    ],
)

doc.add_heading("3.2 Architecture Flow (Bedrock)", level=2)
doc.add_paragraph(
    "User speaks \u2192 WebSocket sends audio \u2192 Amazon Transcribe (streaming) transcribes \u2192 "
    "Bedrock LLM (Claude/Nova) generates response \u2192 Amazon Polly streams audio back \u2192 "
    "WebSocket returns audio to client."
)
doc.add_paragraph(
    "RAG flow: Document uploaded \u2192 Bedrock Knowledge Bases (managed RAG) or manual chunking \u2192 "
    "Titan Embeddings \u2192 stored in OpenSearch Serverless \u2192 at query time, Bedrock retrieves "
    "and injects context automatically."
)

doc.add_heading("3.3 Single API Key / Credential", level=2)
doc.add_paragraph(
    "All AWS services authenticate via IAM roles\u2014no API keys to manage, rotate, or store. "
    "Lambda execution roles grant access to Bedrock, Transcribe, Polly, and OpenSearch. "
    "This eliminates three API keys (OpenAI, Gemini, Voyage) from the configuration."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. COMPARISON: COST
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Comparison: Cost", level=1)

doc.add_heading("4.1 LLM Inference Costs", level=2)
doc.add_paragraph("Per 1 million tokens (input/output), as of March 2026:")
add_table(
    ["Model", "Input / 1M tokens", "Output / 1M tokens", "Notes"],
    [
        ["Gemini 2.5 Flash", "$0.15", "$0.60", "Current primary. Very competitive pricing."],
        ["Gemini 2.5 Pro", "$1.25", "$10.00", "If upgraded for complex reasoning."],
        ["Bedrock \u2014 Claude Sonnet 4", "$3.00", "$15.00", "High quality but 20x more than Gemini Flash."],
        ["Bedrock \u2014 Claude Haiku 4.5", "$0.80", "$4.00", "Fast model, still 5x Gemini Flash."],
        ["Bedrock \u2014 Amazon Nova Pro", "$0.80", "$3.20", "Amazon\u2019s own model. Mid-range."],
        ["Bedrock \u2014 Amazon Nova Lite", "$0.06", "$0.24", "Cheapest Bedrock option. Quality trade-off."],
        ["Bedrock \u2014 Amazon Nova Micro", "$0.035", "$0.14", "Text-only, lowest tier."],
        ["Bedrock \u2014 Llama 4 Scout", "$0.17", "$0.17", "Competitive with Gemini Flash."],
    ],
)

doc.add_heading("4.2 STT (Speech-to-Text) Costs", level=2)
add_table(
    ["Service", "Pricing", "Notes"],
    [
        ["OpenAI Whisper (gpt-4o-transcribe)", "$0.006 / minute", "High accuracy, 100+ languages, real-time capable"],
        ["Amazon Transcribe (real-time)", "$0.024 / minute (first 250K min), $0.015 after", "4x more expensive than Whisper. Custom vocabulary support."],
        ["Amazon Transcribe (batch)", "$0.024 / minute (first 250K min), $0.012 after", "Batch only, not suitable for real-time coaching chat."],
    ],
)

doc.add_heading("4.3 TTS (Text-to-Speech) Costs", level=2)
add_table(
    ["Service", "Pricing", "Notes"],
    [
        ["OpenAI TTS (gpt-4o-mini-tts)", "$0.015 / 1M characters (~$0.10/min)", "Expressive, natural-sounding, instruction-following"],
        ["Gemini 2.5 Flash TTS", "Included in token pricing", "Multimodal output, competitive quality"],
        ["Amazon Polly (Neural)", "$0.016 / 1M characters", "Similar price to OpenAI. Good but less expressive."],
        ["Amazon Polly (Generative)", "$0.030 / 1M characters", "Better quality, 2x neural price. Limited voices."],
        ["Amazon Polly (Standard)", "$0.004 / 1M characters", "Cheapest but robotic. Not suitable for coaching."],
    ],
)

doc.add_heading("4.4 Embeddings Costs", level=2)
add_table(
    ["Service", "Pricing / 1M tokens", "Dimensions", "Notes"],
    [
        ["Google gemini-embedding-001", "Free (Gemini API)", "768", "Currently used. Free tier generous."],
        ["OpenAI text-embedding-3-small", "$0.02 / 1M tokens", "512", "Backup. Very cheap."],
        ["Bedrock \u2014 Titan Embeddings v2", "$0.02 / 1M tokens", "1024", "Same price as OpenAI."],
        ["Bedrock \u2014 Cohere Embed v3", "$0.10 / 1M tokens", "1024", "5x more. Better multilingual."],
        ["Voyage AI", "$0.06 / 1M tokens", "1024", "Configured but unused. Good for code/retrieval."],
    ],
)

doc.add_heading("4.5 Vector Database Costs", level=2)
add_table(
    ["Service", "Monthly Cost", "Notes"],
    [
        ["Milvus Lite (current)", "$0 (self-hosted/embedded)", "Runs in-process. No separate infra cost."],
        ["Milvus (managed \u2014 Zilliz)", "$65+/month", "If externalized."],
        ["OpenSearch Serverless", "$175+/month minimum", "2 OCU minimum even when idle ($0.24/OCU-hr). Expensive for small scale."],
        ["Amazon Aurora pgvector", "$50\u2013$150/month", "If already using Aurora for org data. Cheaper option."],
    ],
)

doc.add_heading("4.6 Cost Summary", level=2)
doc.add_paragraph("Monthly estimated costs for ~1,000 DAU, ~10,000 conversations/month:")
add_table(
    ["Category", "Current Multi-Provider", "AWS Bedrock Standardized", "Delta"],
    [
        ["LLM Inference", "$50\u2013$150 (Gemini Flash)", "$200\u2013$800 (Claude Sonnet/Nova Pro)", "+$150\u2013$650"],
        ["STT", "$30\u2013$60 (OpenAI Whisper)", "$120\u2013$240 (Amazon Transcribe)", "+$90\u2013$180"],
        ["TTS", "$20\u2013$50 (OpenAI + Gemini)", "$30\u2013$80 (Polly Neural/Generative)", "+$10\u2013$30"],
        ["Embeddings", "$0\u2013$5 (Google free tier)", "$5\u2013$15 (Titan)", "+$5\u2013$15"],
        ["Vector DB", "$0 (Milvus Lite embedded)", "$175\u2013$350 (OpenSearch Serverless)", "+$175\u2013$350"],
        ["API Key Management", "3 keys to manage", "0 keys (IAM roles)", "Operational savings"],
        ["TOTAL", "$100\u2013$265/month", "$530\u2013$1,485/month", "+$430\u2013$1,220/month"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("VERDICT: Current multi-provider is significantly cheaper. AWS Bedrock standardization would increase AI costs by 3\u20135x.")
run.bold = True
run.font.color.rgb = RGBColor(180, 0, 0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. AVAILABILITY & RELIABILITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Comparison: Availability & Reliability", level=1)

add_table(
    ["Dimension", "Multi-Provider", "AWS Bedrock"],
    [
        ["SLA", "OpenAI: 99.9% uptime SLA\nGemini: 99.9% SLA (paid tier)\nNo cross-provider SLA", "Bedrock: 99.9% SLA\nTranscribe: 99.9% SLA\nPolly: 99.9% SLA\nAll under single AWS SLA"],
        ["Regional Availability", "OpenAI: US-based, global CDN\nGemini: global endpoints", "Bedrock: Available in us-east-1, us-west-2, eu-west-1, ap-northeast-1, and others. Not all models in all regions."],
        ["Rate Limits", "OpenAI: Tier-based (can be restrictive at scale)\nGemini: 1500 RPM free, 10K+ paid", "Bedrock: On-demand (no pre-provisioning needed) or Provisioned Throughput for guaranteed capacity"],
        ["Outage Risk", "If OpenAI goes down, STT/TTS fails but LLM works\nIf Gemini goes down, LLM fails but STT/TTS works\nPartial resilience through provider diversity", "Single point of failure: if AWS us-east-1 has issues, ALL AI services go down simultaneously\nCan mitigate with multi-region but adds complexity/cost"],
        ["Incident History", "OpenAI: Multiple outages in 2024-2025\nGemini: Fewer reported outages", "AWS: Very reliable historically, but Bedrock is newer and has had capacity issues during launch periods"],
    ],
)

p = doc.add_paragraph()
run = p.add_run(
    "VERDICT: Multi-provider has a natural resilience advantage\u2014if one provider goes down, "
    "others still work. AWS Bedrock consolidates to a single failure domain. However, AWS\u2019s "
    "overall infrastructure reliability is excellent."
)
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. SCALABILITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Comparison: Scalability", level=1)

add_table(
    ["Dimension", "Multi-Provider", "AWS Bedrock"],
    [
        ["Auto-scaling", "Each provider scales independently\nOpenAI/Gemini handle scaling transparently\nNo infrastructure to manage", "Bedrock on-demand scales automatically\nProvisioned Throughput for guaranteed capacity\nTranscribe/Polly auto-scale natively"],
        ["Throughput Limits", "OpenAI: 10K+ RPM on paid tiers\nGemini: 10K RPM on paid tiers\nMust manage rate limits per provider", "Bedrock: Account-level quotas (adjustable)\nDefault: 100\u2013300 RPM per model (varies)\nCan request increases via AWS Support"],
        ["Burst Handling", "Good: API providers absorb bursts\nNo pre-provisioning needed", "Good: On-demand absorbs bursts\nGreat: Provisioned Throughput eliminates throttling\nAdditional cost for provisioned"],
        ["Geographic Scaling", "Global endpoints by default\nNo region-specific deployment needed", "Must deploy to specific AWS regions\nCross-region adds latency + cost\nBut aligns with existing AWS infra"],
        ["Cost at Scale", "Linear: pay per token/minute\nGemini Flash stays cheap at scale\nNo minimum commitments", "Linear: pay per token/minute\nProvisioned Throughput offers discounts at high volume\nOpenSearch Serverless has high minimums"],
    ],
)

p = doc.add_paragraph()
run = p.add_run(
    "VERDICT: Comparable. Both approaches scale well. Multi-provider has slightly better default rate limits. "
    "Bedrock has the advantage of Provisioned Throughput for predictable, high-volume workloads."
)
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. FLEXIBILITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Comparison: Flexibility", level=1)

add_table(
    ["Dimension", "Multi-Provider", "AWS Bedrock"],
    [
        ["Model Selection", "BEST-OF-BREED: Pick the best model for each task\n\u2022 Gemini Flash for cheap, fast reasoning\n\u2022 OpenAI Whisper for best-in-class STT\n\u2022 OpenAI TTS for expressive speech\n\u2022 Can add Claude, Llama, Mistral anytime", "CURATED MARKETPLACE: Access to Claude, Llama, Mistral, Cohere, Amazon Nova, Titan via single API\n\u2022 Cannot use Gemini (not on Bedrock)\n\u2022 Cannot use OpenAI models (not on Bedrock)\n\u2022 TTS limited to Amazon Polly (less expressive)\n\u2022 STT limited to Amazon Transcribe"],
        ["Model Switching", "Requires changing SDK calls per provider\nDifferent APIs, different response formats\nLangChain abstracts some of this", "Easy model switching via model ID parameter\nUnified Bedrock Converse API\nAll models share same request/response format"],
        ["Fine-tuning", "OpenAI: Custom fine-tuning available\nGemini: Fine-tuning available\nFull control over training data", "Bedrock Custom Models: Fine-tune Titan, Llama, Cohere\nCannot fine-tune Claude on Bedrock\nMore limited than direct provider access"],
        ["Prompt Engineering", "Each model has unique prompt formats\nMust optimize per provider\nMore work but better results", "Unified Converse API normalizes this\nLess optimization needed\nBut may sacrifice model-specific tricks"],
        ["Custom Voices (TTS)", "OpenAI: Instruction-following TTS (dynamic tone/style)\nGemini: Multimodal TTS with emotion control", "Polly: Fixed neural voices, limited customization\nPolly Brand Voice (enterprise, expensive)\nLess expressive than OpenAI/Gemini TTS"],
    ],
)

p = doc.add_paragraph()
run = p.add_run(
    "VERDICT: Multi-provider wins decisively on flexibility. The ability to use Gemini\u2019s cost-effective "
    "reasoning + OpenAI\u2019s superior TTS/STT is a significant advantage. Bedrock cannot access Gemini or "
    "OpenAI models, and its TTS (Polly) and STT (Transcribe) are inferior for a coaching platform that "
    "depends on natural, expressive voice interaction."
)
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. MAINTAINABILITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Comparison: Maintainability", level=1)

add_table(
    ["Dimension", "Multi-Provider", "AWS Bedrock"],
    [
        ["API Key Management", "3 API keys to manage, rotate, and secure\n(OpenAI, Gemini, Voyage)\nMust store in environment variables or secrets manager\nKey rotation is manual", "0 API keys \u2014 IAM roles handle all auth\nAutomatic credential rotation\nNo secrets to leak\nClear advantage for security"],
        ["SDK Management", "5+ SDKs to maintain and update:\nopenai, google-genai, google-generativeai, langchain-openai, langchain-google-genai\nBreaking changes across providers", "1\u20132 SDKs: boto3 (AWS SDK) + optionally langchain-aws\nSingle SDK updates\nAWS maintains backward compatibility"],
        ["Error Handling", "Different error formats per provider\nMust handle OpenAI errors, Google errors separately\nRetry logic differs per provider", "Unified error format (AWS exceptions)\nConsistent retry with AWS SDK\nSingle error handling pattern"],
        ["Monitoring", "Must monitor 3+ provider dashboards\nOpenAI dashboard, Google Cloud console, custom metrics\nNo unified view", "Single CloudWatch dashboard\nAll AI metrics in one place\nAlarms, X-Ray tracing integrated\nClear advantage for operations"],
        ["Billing", "3 separate bills (OpenAI, Google, Voyage)\nDifferent billing cycles, currencies\nHarder to track total AI spend", "Single AWS bill\nCost Explorer for all AI services\nBudget alerts unified"],
        ["Dependency Complexity", "8+ AI-related packages in pyproject.toml\nVersion conflicts possible\nLarger deployment artifact", "2\u20133 AI-related packages\nSmaller dependency tree\nFewer version conflicts"],
    ],
)

p = doc.add_paragraph()
run = p.add_run(
    "VERDICT: AWS Bedrock wins on maintainability. Unified auth (IAM), single SDK, unified monitoring, "
    "and single billing are significant operational advantages. The multi-provider approach carries "
    "real maintenance overhead with 3+ API keys and 5+ SDKs."
)
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. EXTENSIBILITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Comparison: Extensibility", level=1)

add_table(
    ["Dimension", "Multi-Provider", "AWS Bedrock"],
    [
        ["Adding New AI Features", "Can adopt any new model from any provider immediately\nE.g., if OpenAI releases GPT-5, integrate same day\nIf Anthropic releases Claude 5, add directly", "New models appear on Bedrock weeks/months after provider launch\nCannot use models not on Bedrock\nAmazon Nova models available immediately"],
        ["RAG / Knowledge Bases", "Manual: LangChain + Milvus, full control\nCan customize chunking, retrieval, reranking\nMore work but more flexible", "Bedrock Knowledge Bases: fully managed RAG\nAutomatic chunking, embedding, retrieval\nLess control but much less code to maintain"],
        ["Agents & Workflows", "Manual: Custom Python agent logic\nFull control over agent behavior\nMore code to maintain", "Bedrock Agents: managed agent framework\nAction groups, knowledge bases, guardrails\nLess custom code, but less flexible"],
        ["Guardrails & Safety", "Must implement content filtering manually\nDifferent moderation APIs per provider\nOpenAI Moderation API is good", "Bedrock Guardrails: managed content filtering\nTopic denial, PII detection, word filters\nApplied uniformly across all models"],
        ["Multi-Modal", "Gemini: native multimodal (text, image, audio, video)\nOpenAI: GPT-4V for vision + TTS/STT\nBest-in-class multimodal", "Bedrock: Claude 4 supports vision\nBut no native audio/video input via Bedrock\nMust use Transcribe/Polly separately"],
        ["Custom Model Hosting", "Not applicable (using hosted APIs)", "Bedrock Custom Model Import: bring your own model\nSageMaker for fully custom training\nUseful if IG trains proprietary coaching models"],
    ],
)

p = doc.add_paragraph()
run = p.add_run(
    "VERDICT: Multi-provider wins for cutting-edge model access and multimodal capabilities. "
    "AWS Bedrock wins for managed services (Knowledge Bases, Agents, Guardrails) that reduce custom code. "
    "For a coaching platform prioritizing voice quality and model freshness, multi-provider has the edge."
)
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. SIDE-BY-SIDE SUMMARY MATRIX
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Side-by-Side Summary Matrix", level=1)

add_table(
    ["Dimension", "Multi-Provider", "AWS Bedrock", "Winner"],
    [
        ["Cost", "$100\u2013$265/mo", "$530\u2013$1,485/mo", "Multi-Provider (3\u20135x cheaper)"],
        ["Availability", "Natural redundancy across providers", "Single AWS failure domain, 99.9% SLA", "Multi-Provider (slight edge)"],
        ["Scalability", "Good auto-scaling, higher default rate limits", "Good auto-scaling, Provisioned Throughput option", "Tie"],
        ["Flexibility", "Best-of-breed model selection, superior TTS/STT", "Curated marketplace, no Gemini/OpenAI models", "Multi-Provider (decisive)"],
        ["Maintainability", "3+ API keys, 5+ SDKs, 3 dashboards", "0 API keys, 1 SDK, unified CloudWatch", "AWS Bedrock (decisive)"],
        ["Extensibility", "Immediate access to latest models", "Managed RAG, Agents, Guardrails", "Tie (different strengths)"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Overall Score: Multi-Provider wins 3 categories, Bedrock wins 1, 2 ties.")
run.bold = True
run.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. PROS AND CONS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Pros and Cons", level=1)

doc.add_heading("11.1 Pros of Staying Multi-Provider", level=2)
bullet("Significantly lower cost: Gemini Flash is 5\u201320x cheaper than comparable Bedrock models for LLM inference.")
bullet("Best-of-breed quality: OpenAI\u2019s TTS (gpt-4o-mini-tts) produces the most expressive, natural coaching voice. Amazon Polly cannot match this for a coaching use case where voice warmth matters.")
bullet("Best-of-breed STT: OpenAI Whisper is the gold standard for speech recognition. Amazon Transcribe is good but costs 4x more and has lower accuracy for conversational speech.")
bullet("Provider diversity = resilience: If Google has an outage, TTS/STT still works via OpenAI. If OpenAI has an outage, reasoning still works via Gemini.")
bullet("Immediate access to innovation: When a new model launches (GPT-5, Claude 5, Gemini 3 Pro), you can adopt it immediately without waiting for Bedrock marketplace listing.")
bullet("Free embeddings: Google\u2019s gemini-embedding-001 is currently free on the Gemini API, saving $5\u2013$15/month vs Bedrock Titan Embeddings.")
bullet("Milvus Lite is zero cost: Running embedded in-process. OpenSearch Serverless has a $175/month minimum.")
bullet("LangChain abstraction already in place: The codebase already uses LangChain wrappers, making provider swaps relatively easy without a full rewrite.")

doc.add_heading("11.2 Cons of Staying Multi-Provider", level=2)
bullet("3 API keys to manage, rotate, and secure. Key leakage risk increases with more providers.")
bullet("5+ SDK dependencies to maintain. Breaking changes across providers create update burden.")
bullet("3 separate billing dashboards. Harder to track and forecast total AI costs.")
bullet("Different error handling patterns per provider. Increases code complexity and testing surface.")
bullet("No unified monitoring. Must check OpenAI dashboard, Google Cloud console, and custom metrics separately.")
bullet("Vendor relationship management. Must maintain accounts, billing, and support with 3+ providers.")
bullet("Provider deprecation risk. If a provider deprecates a model (e.g., OpenAI sunset gpt-4o-transcribe), migration is forced and urgent.")

doc.add_heading("11.3 Pros of AWS Bedrock Standardization", level=2)
bullet("Zero API keys: IAM roles handle all authentication. No secrets to rotate or risk leaking.")
bullet("Single SDK (boto3): One dependency for all AI services. Fewer version conflicts, smaller artifact.")
bullet("Unified monitoring: All AI metrics in CloudWatch. Single dashboard, unified alerts, X-Ray tracing.")
bullet("Single bill: AWS Cost Explorer shows all AI spend in one place. Easier budgeting and forecasting.")
bullet("AWS ecosystem integration: Native connections to Lambda, S3, DynamoDB, EventBridge. No custom glue code.")
bullet("Managed RAG: Bedrock Knowledge Bases automates document chunking, embedding, and retrieval. Reduces custom RAG code significantly.")
bullet("Guardrails: Built-in content filtering, PII detection, and topic denial. Applied uniformly across models.")
bullet("Provisioned Throughput: Guaranteed capacity for high-traffic periods. Predictable latency.")
bullet("HIPAA/SOC2 compliance: AWS Bedrock is HIPAA-eligible and SOC2 certified. Important if IG handles sensitive coaching data.")

doc.add_heading("11.4 Cons of AWS Bedrock Standardization", level=2)
bullet("3\u20135x higher monthly cost ($530\u2013$1,485 vs $100\u2013$265). This is the dealbreaker at current scale.")
bullet("Cannot use Gemini models. Google\u2019s Gemini is not available on Bedrock. Losing access to the cheapest high-quality LLM.")
bullet("Cannot use OpenAI models. GPT-4o, Whisper, and OpenAI TTS are not on Bedrock. Losing best-in-class STT and TTS.")
bullet("Amazon Polly TTS is less expressive than OpenAI TTS. For an AI coaching platform where voice warmth and naturalness drive user engagement, this is a meaningful quality regression.")
bullet("Amazon Transcribe costs 4x more than Whisper for similar accuracy.")
bullet("OpenSearch Serverless has a $175/month minimum. Milvus Lite is free.")
bullet("Single failure domain. All AI services depend on AWS. A regional AWS outage takes down everything simultaneously.")
bullet("Model availability lag. New models from Anthropic, Meta, etc. appear on Bedrock weeks to months after direct availability.")
bullet("Deep vendor lock-in. Migrating away from Bedrock+Transcribe+Polly+OpenSearch is far more work than switching API keys.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. BENEFITS AND RISKS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Benefits and Risks", level=1)

doc.add_heading("12.1 Benefits of Standardizing on Bedrock", level=2)
bullet("Operational simplicity: One platform for all AI, one monitoring dashboard, one bill, one set of IAM policies.")
bullet("Security posture: No API keys in environment variables. IAM roles are the gold standard for AWS credential management.")
bullet("Compliance readiness: Bedrock is HIPAA-eligible, SOC2 certified, and offers data residency guarantees. Critical if IG expands to enterprise clients.")
bullet("Developer experience: New developers only need to learn the AWS SDK, not 3+ provider APIs.")
bullet("Managed RAG eliminates: Milvus deployment, embedding pipeline code, chunking logic, and retrieval tuning.")
bullet("Guardrails protect against: Prompt injection, PII leakage, off-topic responses\u2014without custom code.")
bullet("Future-proofing for AWS migration: If the backend moves to Lambda (per the architecture plan), Bedrock is the natural fit.")

doc.add_heading("12.2 Risk Matrix", level=2)
add_table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Voice quality regression (Polly < OpenAI TTS)", "High", "High", "A/B test with users before committing. Keep OpenAI TTS as fallback."],
        ["Cost increase (3\u20135x)", "Certain", "High", "Use Amazon Nova Lite/Micro for simple tasks. Optimize prompt length. Monitor token usage."],
        ["STT accuracy regression (Transcribe < Whisper)", "Medium", "Medium", "Benchmark with real coaching audio. Transcribe has improved significantly."],
        ["Model quality regression (if Gemini Flash > Nova)", "Medium", "High", "Use Claude Sonnet on Bedrock for quality. Accept higher cost or use Llama 4."],
        ["Migration effort disrupts feature development", "High", "Medium", "Phase migration over sprints. Never migrate all services at once."],
        ["Vendor lock-in to AWS", "Certain", "Medium", "Already on AWS (S3, CloudFront). Bedrock deepens lock-in but AWS is stable."],
        ["OpenSearch Serverless cost floor ($175/mo)", "Certain", "Low", "Use Aurora pgvector instead if already on Aurora for org data."],
        ["Bedrock model deprecation", "Low", "Medium", "AWS provides 6+ month deprecation notices. Multiple model choices available."],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 13. MIGRATION EFFORT ESTIMATE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Migration Effort Estimate", level=1)

doc.add_paragraph("If you decided to migrate to AWS Bedrock, here is the estimated effort:")

add_table(
    ["Component", "Current", "Target", "Effort", "Risk"],
    [
        ["LLM Reasoning", "Gemini SDK \u2192 generate_and_stream_gemini_res()", "Bedrock Converse API (boto3)", "2\u20133 weeks", "Medium \u2014 different streaming API"],
        ["Query Refinement", "Gemini structured output", "Bedrock Converse with tool_use", "1 week", "Low"],
        ["STT", "OpenAI audio_transcription()", "Amazon Transcribe streaming SDK", "2 weeks", "High \u2014 different streaming protocol"],
        ["TTS (streaming)", "OpenAI stream_sentence_audio()", "Amazon Polly SynthesizeSpeech", "2 weeks", "High \u2014 Polly chunking differs from OpenAI"],
        ["Embeddings", "LangChain GoogleGenerativeAIEmbeddings", "LangChain BedrockEmbeddings", "2\u20133 days", "Low \u2014 LangChain abstracts this"],
        ["Vector DB", "Milvus Lite (embedded)", "OpenSearch Serverless or Aurora pgvector", "1\u20132 weeks", "Medium \u2014 data migration, index rebuild"],
        ["RAG Pipeline", "Manual LangChain chain", "Bedrock Knowledge Bases (optional)", "1\u20132 weeks", "Medium \u2014 less control over chunking"],
        ["Auth/Credentials", "3 API keys in env vars", "IAM roles on Lambda", "2\u20133 days", "Low"],
        ["Testing & Validation", "\u2014", "Quality benchmarks, A/B testing", "2\u20133 weeks", "High \u2014 must validate voice quality"],
        ["TOTAL", "\u2014", "\u2014", "10\u201314 weeks (solo dev)", "\u2014"],
    ],
)

doc.add_paragraph(
    "Development cost estimate: $30,000\u2013$84,000 (at $75\u2013$150/hr, 10\u201314 weeks, 1\u20132 developers)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 14. OPINION & RECOMMENDATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Opinion & Recommendation", level=1)

doc.add_heading("14.1 Is the Change Worth the Effort?", level=2)

p = doc.add_paragraph()
run = p.add_run("Short answer: No, not as a full standardization. A selective hybrid approach is recommended instead.")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 100, 0)

doc.add_paragraph(
    "The current multi-provider strategy is working well for Inspire Genius. The codebase is cleanly "
    "architected with LangChain abstractions that already make provider switching relatively easy. "
    "Gemini Flash provides exceptional value for LLM inference at a fraction of Bedrock model costs. "
    "OpenAI\u2019s TTS produces the warm, expressive coaching voice that is central to the platform\u2019s "
    "user experience. And OpenAI Whisper remains the most cost-effective, highest-quality STT option."
)

doc.add_heading("14.2 Why Full Bedrock Standardization is Not Recommended", level=2)
bullet("Cost increase is prohibitive: 3\u20135x higher monthly AI costs ($430\u2013$1,220 more per month) with no revenue increase to offset it.")
bullet("Voice quality regression: Amazon Polly cannot match OpenAI\u2019s instruction-following TTS for coaching. Voice quality is a core differentiator for an AI coaching platform.")
bullet("Losing Gemini: The cheapest high-quality LLM is not available on Bedrock. No Bedrock model matches Gemini Flash\u2019s price/performance ratio.")
bullet("Migration effort: 10\u201314 weeks of development time that could be spent on features, bug fixes, or the microservices migration.")
bullet("Deepened vendor lock-in: Moving to Bedrock + Transcribe + Polly + OpenSearch makes future changes much harder.")

doc.add_heading("14.3 Recommended Hybrid Approach", level=2)
doc.add_paragraph(
    "Instead of full standardization, adopt a selective hybrid strategy that uses AWS services where "
    "they add clear value while keeping best-of-breed providers where quality matters:"
)

add_table(
    ["Function", "Recommendation", "Rationale"],
    [
        ["LLM Reasoning", "KEEP Gemini Flash", "5\u201320x cheaper than Bedrock alternatives. Quality is excellent."],
        ["STT", "KEEP OpenAI Whisper", "4x cheaper than Transcribe. Best accuracy. Already integrated."],
        ["TTS", "KEEP OpenAI TTS", "Most expressive, natural coaching voice. Core UX differentiator."],
        ["Embeddings", "KEEP Google (free) or SWITCH to Bedrock Titan", "Either works. Titan is $0.02/1M tokens, marginal cost. LangChain makes swapping trivial."],
        ["Vector DB", "KEEP Milvus Lite for now", "Free and works. Switch to Aurora pgvector if/when migrating to Aurora for org data."],
        ["Guardrails", "ADD Bedrock Guardrails as a middleware layer", "Use Bedrock Guardrails to filter LLM input/output without changing the LLM provider. Adds safety without lock-in."],
        ["RAG", "EVALUATE Bedrock Knowledge Bases", "If the manual RAG pipeline becomes a maintenance burden, Bedrock KB can reduce code. But only if the quality meets needs."],
        ["Credentials", "MIGRATE API keys to AWS Secrets Manager", "Keep using provider APIs but store keys in Secrets Manager instead of env vars. Get IAM-level security without changing providers."],
    ],
)

doc.add_heading("14.4 When to Reconsider Full Bedrock", level=2)
doc.add_paragraph("Revisit full Bedrock standardization if any of these conditions become true:")
bullet("Scale exceeds 10,000 DAU and Provisioned Throughput economics make Bedrock cheaper than per-token API pricing.")
bullet("Amazon Polly launches a \u201cGenerative Engine 2.0\u201d that matches OpenAI TTS quality for coaching-style voices.")
bullet("Gemini is added to the Bedrock marketplace (unlikely but possible if Google and AWS partner).")
bullet("Enterprise clients require HIPAA/FedRAMP compliance and insist on all AI processing within AWS boundaries.")
bullet("The team grows to 5+ engineers and operational overhead of multi-provider becomes a genuine drag on velocity.")

doc.add_heading("14.5 Final Recommendation Summary", level=2)
add_table(
    ["Action", "Priority", "Effort", "Impact"],
    [
        ["Keep Gemini Flash for LLM", "Now", "None", "Save $150\u2013$650/month vs Bedrock"],
        ["Keep OpenAI for STT/TTS", "Now", "None", "Best voice quality for coaching"],
        ["Move API keys to Secrets Manager", "Next sprint", "1\u20132 days", "Better security posture"],
        ["Evaluate Bedrock Guardrails", "Q2 2026", "1 week", "Add content safety without provider change"],
        ["Evaluate Bedrock Knowledge Bases", "Q3 2026", "2 weeks", "Potentially reduce RAG maintenance"],
        ["Full Bedrock standardization", "Not recommended at current scale", "10\u201314 weeks", "Negative ROI at <10K DAU"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    "The multi-provider strategy is not a liability\u2014it\u2019s a competitive advantage. "
    "It gives Inspire Genius the best models at the best prices. The operational overhead "
    "of managing 3 API keys and 5 SDKs is real but manageable, and far less costly than "
    "the 3\u20135x price increase and voice quality regression that full Bedrock standardization would bring."
)
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("\u2014 End of Document \u2014")
run.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── SAVE ────────────────────────────────────────────────────────────────
output_dir = "/Users/williambrown/Dropbox/AES Material/Inspire-X/New IG Projects/Local_IG-App_UI/inspire-genius-frontend"
output_path = os.path.join(output_dir, "Inspire_Genius_LLM_Strategy_Comparison.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")