"""Generate Phased Correlated Implementation Plan Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0, 51, 102)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def bold_para(text, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Inspire Genius\nPhased Implementation Plan")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    "Correlated Roadmap Across All Workstreams\n"
    "Bug Stabilization \u2022 New UI Adoption \u2022 PWA \u2022 Microservices \u2022 LLM Strategy"
)
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0, 128, 128)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Prepared for: ").bold = True
meta.add_run("AES / Inspire Genius Team\n")
meta.add_run("Date: ").bold = True
meta.add_run("March 14, 2026\n")
meta.add_run("Classification: ").bold = True
meta.add_run("Internal \u2014 Strategic Roadmap\n\n")
meta.add_run("Related Documents:\n").bold = True
meta.add_run("  \u2022 Inspire_Genius_Architecture_Plan.docx\n")
meta.add_run("  \u2022 Inspire_Genius_LLM_Strategy_Comparison.docx\n")
meta.add_run("  \u2022 Inspire_Genius_New_UI_Adoption_Plan.docx")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc = [
    "1. Executive Summary",
    "2. Workstream Overview",
    "3. Dependency Map",
    "4. Phase 0: Stabilize (Weeks 1\u20134)",
    "5. Phase 1: Foundation (Weeks 5\u20138)",
    "6. Phase 2: Core User Experience (Weeks 9\u201314)",
    "7. Phase 3: Management & Organization (Weeks 15\u201322)",
    "8. Phase 4: Practitioner & Distributor (Weeks 23\u201328)",
    "9. Phase 5: Platform Intelligence (Weeks 29\u201334)",
    "10. Phase 6: Infrastructure Modernization (Weeks 35\u201344)",
    "11. Phase 7: Super Admin & Polish (Weeks 45\u201348)",
    "12. Consolidated Timeline",
    "13. Resource Requirements",
    "14. Go/No-Go Decision Gates",
    "15. Risk Register",
    "16. Success Metrics",
]
for item in toc:
    doc.add_paragraph(item, style="List Number")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This document presents a unified, phased implementation roadmap that correlates all "
    "planned workstreams for the Inspire Genius platform into a single sequenced plan. "
    "Three previous planning documents each addressed a distinct initiative:"
)
bullet("Architecture Plan \u2014 PWA conversion, microservices decomposition, event-driven serverless migration")
bullet("LLM Strategy Comparison \u2014 multi-provider vs AWS Bedrock standardization (recommendation: keep multi-provider, adopt select Bedrock services)")
bullet("New UI Adoption Plan \u2014 role-based single-page PWA with 5 user sections (User, Manager, Company Admin, Practitioner, Distributor)")

doc.add_paragraph(
    "This plan sequences all three initiatives into 8 phases spanning approximately 48 weeks "
    "(12 months). Each phase has clear entry criteria, deliverables, exit criteria, and "
    "dependencies on other phases. The plan accounts for the current reality: the team is "
    "actively fixing bugs from a recent server migration."
)

doc.add_paragraph("Key principles:")
bullet("Never stop fixing bugs \u2014 20\u201330% of each sprint is reserved for bug fixes throughout all phases")
bullet("Ship incrementally \u2014 every phase delivers standalone user-visible value")
bullet("Frontend leads, backend follows \u2014 new UI pages can be built with mock data while backend APIs are developed")
bullet("Strangler Fig for backend \u2014 new microservices are deployed behind API Gateway alongside the monolith")
bullet("LLM strategy stays multi-provider \u2014 selectively adopt Bedrock Guardrails and Secrets Manager, not a full migration")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. WORKSTREAM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Workstream Overview", level=1)

doc.add_paragraph("Five concurrent workstreams are woven through the phases:")

add_table(
    ["Workstream", "Color Code", "Description", "Source Document"],
    [
        ["WS-A: Bug Stabilization", "RED", "Fix critical/high bugs from server migration. Ongoing throughout.", "N/A (current work)"],
        ["WS-B: New UI Frontend", "TEAL", "Build unified role-based PWA dashboard with 5 sections + 35 new routes.", "New_UI_Adoption_Plan.docx"],
        ["WS-C: Backend APIs", "BLUE", "Build 15 new API domains to support new UI features (PRISM, Goals, Credits, etc.).", "New_UI_Adoption_Plan.docx"],
        ["WS-D: Infrastructure", "ORANGE", "PWA conversion, IaC (CDK), API Gateway, microservices extraction, monitoring.", "Architecture_Plan.docx"],
        ["WS-E: AI/LLM Optimization", "PURPLE", "API key security, Bedrock Guardrails evaluation, LLM cost optimization.", "LLM_Strategy_Comparison.docx"],
    ],
)

doc.add_heading("2.1 Workstream Correlation Rules", level=2)
bullet("WS-B (Frontend) can proceed with mock data even when WS-C (Backend) is not ready")
bullet("WS-D (Infrastructure) runs independently of WS-B/WS-C until API Gateway proxy phase")
bullet("WS-E (AI/LLM) is lightweight and fits into gaps between other work")
bullet("WS-A (Bugs) is ongoing and never blocks other workstreams")
bullet("When WS-C delivers an API, WS-B immediately integrates (replacing mock data)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. DEPENDENCY MAP
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Dependency Map", level=1)

doc.add_paragraph("Critical dependencies between deliverables:")

add_table(
    ["Deliverable", "Depends On", "Blocks"],
    [
        ["Unified Layout (WS-B)", "Multi-role Auth (WS-B)", "All new UI sections"],
        ["Multi-role Auth (WS-B)", "Backend roles[] API change (WS-C)", "Unified Layout"],
        ["PRISM Report Page (WS-B)", "PRISM Assessment API (WS-C)", "Manager Team, Company Dashboard"],
        ["Goals Page (WS-B)", "Goals API (WS-C)", "Nothing (independent)"],
        ["Manager Section (WS-B)", "Team API + PRISM API (WS-C)", "Nothing"],
        ["Practitioner Section (WS-B)", "Client Mgmt API + Credits API (WS-C)", "Nothing"],
        ["Distributor Section (WS-B)", "Credits API + Distributor API (WS-C)", "Nothing"],
        ["Cost Dashboard (WS-B)", "Cost Analytics API (WS-C)", "Nothing"],
        ["API Gateway Proxy (WS-D)", "CDK Foundation (WS-D)", "Microservice extraction"],
        ["Auth Microservice (WS-D)", "API Gateway Proxy (WS-D)", "Remaining microservices"],
        ["Bedrock Guardrails (WS-E)", "Nothing", "Nothing (additive)"],
        ["Secrets Manager Migration (WS-E)", "Nothing", "Nothing (additive)"],
        ["PWA Service Worker (WS-D)", "Nothing", "Nothing (additive)"],
    ],
)

doc.add_heading("3.1 Critical Path", level=2)
doc.add_paragraph(
    "The longest dependency chain determines the minimum timeline:"
)
doc.add_paragraph(
    "Stabilize Bugs \u2192 Multi-role Auth API \u2192 Unified Layout \u2192 User Section + PRISM API "
    "\u2192 Manager Section + Team API \u2192 Company Section \u2192 Practitioner Section + Credits API "
    "\u2192 Distributor Section \u2192 Super Admin Migration \u2192 Infrastructure Modernization"
)
doc.add_paragraph(
    "Estimated critical path duration: 40\u201348 weeks (10\u201312 months)"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. PHASE 0: STABILIZE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Phase 0: Stabilize (Weeks 1\u20134)", level=1)

bold_para("Goal: Reach a stable baseline where core user flows work reliably.", 12, RGBColor(0, 100, 0))

doc.add_heading("4.1 Entry Criteria", level=2)
bullet("Application is deployed and accessible")
bullet("Known bug list exists (even if informal)")

doc.add_heading("4.2 Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["0.1", "Critical bug fixes", "WS-A", "Fix all login, auth, and data-corruption bugs. Users must be able to log in, chat with coaches, upload documents."],
        ["0.2", "High-severity bug fixes", "WS-A", "Fix UI rendering issues, broken navigation, WebSocket disconnections, and token refresh failures."],
        ["0.3", "Bug triage & prioritization", "WS-A", "Categorize all known bugs as Critical, High, Medium, Low. Create tracking system."],
        ["0.4", "Environment documentation", "WS-D", "Document all environment variables, deployment steps, AWS resources. Create .env.example."],
        ["0.5", "API key audit", "WS-E", "Inventory all API keys (OpenAI, Gemini, Voyage). Document rotation schedule. Verify secure storage."],
    ],
)

doc.add_heading("4.3 Exit Criteria (Go/No-Go for Phase 1)", level=2)
bullet("Users can log in and maintain sessions without errors")
bullet("Coach chat (text + voice) works end-to-end")
bullet("Document upload/download works")
bullet("No data-corruption bugs remain")
bullet("All Critical bugs resolved; High bugs have workarounds or fixes")

doc.add_heading("4.4 Effort", level=2)
doc.add_paragraph("2\u20134 weeks, 1\u20132 developers, $6,000\u2013$24,000")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. PHASE 1: FOUNDATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Phase 1: Foundation (Weeks 5\u20138)", level=1)

bold_para("Goal: Build the architectural foundation for the new UI and infrastructure.", 12, RGBColor(0, 100, 0))

doc.add_heading("5.1 Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details", "Dependencies"],
    [
        ["1.1", "Multi-role Auth API", "WS-C", "Backend returns roles as string[] instead of single string. Update /v1/me and /v1/login response payloads.", "Phase 0 complete"],
        ["1.2", "Multi-role AuthContext", "WS-B", "Update AuthContext to store roles[]. Update secureStorage, STORAGE_KEYS, hasAccess(). Update ProtectedRoute.", "1.1"],
        ["1.3", "Unified Layout", "WS-B", "Build UnifiedLayout.tsx, UnifiedSidebar.tsx, SidebarNavGroup.tsx, BreadcrumbHeader.tsx. Replace UserLayout + SuperAdminLayout.", "1.2"],
        ["1.4", "Route Restructuring", "WS-B", "Move /dashboard to /chat. Create /dashboard as new user home. Add route stubs for all 35 new routes.", "1.3"],
        ["1.5", "Expanded Role Constants", "WS-B", "Add manager, practitioner, distributor to ROLES enum. Update ROLE_PERMISSIONS. Add ROLE_LABELS.", "None"],
        ["1.6", "PWA Manifest + Icons", "WS-D", "Create manifest.webmanifest, add PWA icons (192, 512, maskable), update index.html with manifest link and meta tags.", "None"],
        ["1.7", "API Keys to Secrets Manager", "WS-E", "Migrate OPENAI_API_KEY, GEMINI_API_KEY from env vars to AWS Secrets Manager. Update backend config to read from Secrets Manager.", "None"],
        ["1.8", "MetricCard Component", "WS-B", "Build reusable MetricCard.tsx (value, label, trend, icon). Will be used across all dashboard sections.", "None"],
    ],
)

doc.add_heading("5.2 Parallel Work Allocation", level=2)
add_table(
    ["Developer 1 (Frontend)", "Developer 2 (Backend/Infra)"],
    [
        ["1.2 Multi-role AuthContext", "1.1 Multi-role Auth API"],
        ["1.3 Unified Layout", "1.6 PWA Manifest + Icons"],
        ["1.4 Route Restructuring", "1.7 Secrets Manager Migration"],
        ["1.5 Role Constants + 1.8 MetricCard", "Environment docs"],
    ],
)

doc.add_heading("5.3 Exit Criteria", level=2)
bullet("Unified sidebar renders correct nav groups based on user roles[]")
bullet("Users with multiple roles see all applicable sections")
bullet("All existing functionality still works (no regression)")
bullet("manifest.webmanifest passes PWA audit in Chrome DevTools")
bullet("API keys retrieved from Secrets Manager, not env vars")

doc.add_heading("5.4 Effort", level=2)
doc.add_paragraph("4 weeks, 2 developers, $24,000\u2013$48,000")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. PHASE 2: CORE USER EXPERIENCE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Phase 2: Core User Experience (Weeks 9\u201314)", level=1)

bold_para("Goal: Deliver the new User (MAIN) section and PRISM foundation.", 12, RGBColor(0, 100, 0))

doc.add_heading("6.1 Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details", "Dependencies"],
    [
        ["2.1", "User Dashboard Redesign", "WS-B", "Redesign Home.tsx as /dashboard: stat cards (Active Goals, Sessions, PRISM Score, Documents), PRISM profile summary, activity feed, quick actions.", "1.3 Layout"],
        ["2.2", "Chat Hub Move", "WS-B", "Move coach chat to /chat route. Reuse existing ChatWindow, ChatWindowChatTab, MessageFeedback. Update navigation.", "1.4 Routes"],
        ["2.3", "PRISM Assessment API", "WS-C", "Build PRISM service endpoints: POST /v1/prism/assessment (submit), GET /v1/prism/report/{userId} (results), GET /v1/prism/team/{teamId} (team data).", "None"],
        ["2.4", "PRISM Visualization Components", "WS-B", "Build PrismColorWheel.tsx, PrismScoreCard.tsx (Recharts). Red=Drive, Blue=Compliance, Green=Steadiness, Yellow=Influence.", "2.3 API"],
        ["2.5", "PRISM Report Page", "WS-B", "Build /prism-report page: PRISM wheel, behavioral scores, strengths, development areas, communication style, compatibility insights.", "2.3, 2.4"],
        ["2.6", "Goals API", "WS-C", "Build Goals service: CRUD /v1/goals, GET /v1/goals/stats. Types: GoalEntry, GoalStatus (active, completed, overdue).", "None"],
        ["2.7", "Goals Page", "WS-B", "Build /goals page: goal cards with progress bars, Set New Goal modal (React Hook Form + Zod), filter by status, metrics.", "2.6 API"],
        ["2.8", "Service Worker", "WS-D", "Add vite-plugin-pwa. Configure Workbox caching strategies. NetworkFirst for API, CacheFirst for static assets. Offline banner.", "1.6 Manifest"],
        ["2.9", "Install Prompt Hook", "WS-D", "Build useInstallPrompt.ts, useOnlineStatus.ts, OfflineBanner.tsx. Show install banner after 3rd visit.", "2.8"],
        ["2.10", "Notifications API", "WS-C", "Build Notifications service: GET /v1/notifications, PATCH /v1/notifications/{id}/read. Types: Notification, NotificationListParams.", "None"],
        ["2.11", "Notifications Page", "WS-B", "Build /notifications page: notification list with read/unread, badge count in sidebar.", "2.10"],
    ],
)

doc.add_heading("6.2 Parallel Work (3 developers)", level=2)
add_table(
    ["Frontend Dev 1", "Frontend Dev 2", "Backend Dev"],
    [
        ["2.1 Dashboard Redesign", "2.4 PRISM Components", "2.3 PRISM API"],
        ["2.2 Chat Hub Move", "2.5 PRISM Report Page", "2.6 Goals API"],
        ["2.7 Goals Page", "2.8 Service Worker + 2.9 Install Prompt", "2.10 Notifications API"],
        ["2.11 Notifications Page", "Bug fixes (20%)", "Bug fixes (20%)"],
    ],
)

doc.add_heading("6.3 Exit Criteria", level=2)
bullet("New /dashboard shows stat cards and PRISM summary")
bullet("Coach chat works at /chat (old /dashboard route redirects)")
bullet("PRISM Report page renders real assessment data")
bullet("Goals page supports full CRUD with filtering")
bullet("PWA installs on Chrome/Edge/Safari; offline banner appears when offline")
bullet("Notifications page shows and marks notifications as read")

doc.add_heading("6.4 Effort", level=2)
doc.add_paragraph("6 weeks, 2\u20133 developers, $36,000\u2013$108,000")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. PHASE 3: MANAGEMENT & ORGANIZATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Phase 3: Management & Organization (Weeks 15\u201322)", level=1)

bold_para("Goal: Deliver Manager and Company Admin sections.", 12, RGBColor(0, 100, 0))

doc.add_heading("7.1 Manager Section Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["3.1", "Team API", "WS-C", "GET /v1/team/members, GET /v1/team/analytics, GET /v1/team/prism-balance. Team member PRISM profiles."],
        ["3.2", "Manager Dashboard Page", "WS-B", "/manager/dashboard: Team Members, Active Teams, PRISM Assessed %, Training Completion %, Team Health chart."],
        ["3.3", "My Team Page", "WS-B", "/manager/team: member table with PRISM status, team balance viz (PrismTeamBalance), compatibility insights."],
        ["3.4", "Hiring API", "WS-C", "CRUD /v1/hiring/positions, /v1/hiring/candidates, /v1/hiring/interviews. Pipeline stages."],
        ["3.5", "Hiring Pipeline Page", "WS-B", "/manager/hiring: KPIs (Open Positions, Candidates, Interviews), pipeline board."],
        ["3.6", "Candidates + Interviews Pages", "WS-B", "/manager/candidates: candidate list with PRISM preview. /manager/interviews: schedule view."],
        ["3.7", "Job DNA API", "WS-C", "CRUD /v1/job-dna/profiles, GET /v1/job-dna/match. PRISM-based job profiling."],
        ["3.8", "Job DNA Page", "WS-B", "/manager/job-dna: build ideal PRISM profile for a role, match candidates against it."],
        ["3.9", "Training API", "WS-C", "CRUD /v1/training/programs, /v1/training/enrollment. Program types, completion tracking."],
        ["3.10", "Training + Leadership + Career Pages", "WS-B", "/manager/training, /manager/leadership, /manager/career: program cards, enrollment, career plans."],
        ["3.11", "Team Building Page", "WS-B", "/manager/team-building: PRISM Team Workshop planner, exercise scheduler."],
    ],
)

doc.add_heading("7.2 Company Admin Section Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["3.12", "Org Chart API", "WS-C", "GET /v1/org-chart/hierarchy. Returns tree structure: CEO > VPs > Directors > Managers > ICs."],
        ["3.13", "Company Dashboard Page", "WS-B", "/company/dashboard: Total Employees, Active Teams, PRISM Assessed %, department comparison charts."],
        ["3.14", "Org Chart Page", "WS-B", "/company/org-chart: interactive hierarchy tree. Click node > profile with PRISM data."],
        ["3.15", "Company Teams Page", "WS-B", "/company/teams: team directory, team cards with PRISM balance, Team Management CRUD."],
        ["3.16", "Company Training Page", "WS-B", "/company/training: company-wide program management, enrollment tracking, department assignment."],
        ["3.17", "Assignments + Leadership Pages", "WS-B", "/company/assignments: coach + PRISM assessment assignment. /company/leadership: org-wide leadership profiles."],
    ],
)

doc.add_heading("7.3 Infrastructure Work (Parallel)", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["3.18", "CDK Project Setup", "WS-D", "Initialize infrastructure/ CDK project. Create network-stack, api-gateway-stack, frontend-stack."],
        ["3.19", "API Gateway as Transparent Proxy", "WS-D", "Deploy HTTP API Gateway. Route all /v1/* to existing monolith. Zero behavior change. Strangler Fig foundation."],
        ["3.20", "Bedrock Guardrails Evaluation", "WS-E", "Evaluate Bedrock Guardrails as middleware: content filtering, PII detection, topic denial. POC with 1 agent."],
    ],
)

doc.add_heading("7.4 Exit Criteria", level=2)
bullet("Manager users see MANAGER nav group with 10 working sub-pages")
bullet("Company Admin users see COMPANY nav group with 6 working sub-pages")
bullet("Org chart renders interactive hierarchy with click-to-profile")
bullet("API Gateway proxy serves all traffic transparently (users notice no change)")
bullet("Bedrock Guardrails POC report delivered with recommendation")

doc.add_heading("7.5 Effort", level=2)
doc.add_paragraph("8 weeks, 2\u20133 developers, $48,000\u2013$144,000")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. PHASE 4: PRACTITIONER & DISTRIBUTOR
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Phase 4: Practitioner & Distributor (Weeks 23\u201328)", level=1)

bold_para("Goal: Deliver Practitioner and Distributor sections with credit system.", 12, RGBColor(0, 100, 0))

doc.add_heading("8.1 Practitioner Section Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["4.1", "Practitioner-Client API", "WS-C", "CRUD /v1/practitioner/clients. Client profiles, session history, PRISM results, satisfaction scores."],
        ["4.2", "Credits API", "WS-C", "GET /v1/credits/balance, POST /v1/credits/allocate, POST /v1/credits/purchase. Server-side ledger with transactions."],
        ["4.3", "Follow-Up API", "WS-C", "CRUD /v1/practitioner/follow-ups. Due dates, status (due today, overdue, upcoming, completed)."],
        ["4.4", "Practitioner Dashboard", "WS-B", "/practitioner/dashboard: Active Clients, Sessions This Month, Credits Remaining, Avg Satisfaction."],
        ["4.5", "My Clients Page", "WS-B", "/practitioner/clients: client list, search, filter, click > detail panel with PRISM, history, notes."],
        ["4.6", "Conversations Page", "WS-B", "/practitioner/conversations: conversation history across clients. Reuse ChatWindow for transcript viewing."],
        ["4.7", "Follow-Up Page", "WS-B", "/practitioner/follow-up: task list with due dates, status filtering, create reminders."],
        ["4.8", "Practitioner Credits Page", "WS-B", "/practitioner/credits: balance display, usage history, consumption per session/client."],
    ],
)

doc.add_heading("8.2 Distributor Section Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["4.9", "Distributor API", "WS-C", "CRUD /v1/distributor/practitioners. Practitioner management, credit allocation, performance metrics."],
        ["4.10", "Distributor Dashboard", "WS-B", "/distributor/dashboard: Total Practitioners, Active Practitioners, Total Credits, Credits Allocated."],
        ["4.11", "Practitioners Page", "WS-B", "/distributor/practitioners: roster table, add/deactivate practitioners, view performance."],
        ["4.12", "Distributor Credits Page", "WS-B", "/distributor/credits: pool management, purchase workflow, transaction history."],
        ["4.13", "Credit Allocation Page", "WS-B", "/distributor/credit-allocation: select practitioners, enter amounts, confirm allocation."],
    ],
)

doc.add_heading("8.3 Messages System (Parallel)", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["4.14", "Messages API", "WS-C", "GET/POST /v1/messages. Internal messaging between users, practitioners, managers."],
        ["4.15", "Messages Page", "WS-B", "/messages: message list, compose, reply. Badge count in sidebar."],
    ],
)

doc.add_heading("8.4 Exit Criteria", level=2)
bullet("Practitioner users see 5 working sub-pages with client management and credit tracking")
bullet("Distributor users see 4 working sub-pages with credit allocation and practitioner management")
bullet("Credit system enforces server-side ledger (frontend is display-only)")
bullet("Messages system works for internal communication")

doc.add_heading("8.5 Effort", level=2)
doc.add_paragraph("6 weeks, 2\u20133 developers, $36,000\u2013$108,000")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. PHASE 5: PLATFORM INTELLIGENCE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Phase 5: Platform Intelligence (Weeks 29\u201334)", level=1)

bold_para("Goal: Deliver Cost Dashboard, Assessment tools, and AI optimizations.", 12, RGBColor(0, 100, 0))

doc.add_heading("9.1 Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["5.1", "Cost Analytics API", "WS-C", "GET /v1/cost/platform (KPIs, cost by week, distribution, agent usage, top orgs, top users). GET /v1/cost/company/{id} (department breakdown). GET /v1/cost/team/{id} (member breakdown). GET /v1/cost/user/{id} (session/token/audio cost breakdown)."],
        ["5.2", "Cost Platform Dashboard", "WS-B", "Platform-level: Total Cost MTD, Active Users, New Users, Cost/User, Sessions. Cost by Week pie chart, Cost Distribution pie, Agent Usage pie, Top Orgs table, Top Users table."],
        ["5.3", "Cost Company Dashboard", "WS-B", "Company-level: Org Cost, Active Users, Cost/User, Sessions. Department Cost Breakdown table (dept, users, sessions, AI cost, total, $/user)."],
        ["5.4", "Cost Team Dashboard", "WS-B", "Team-level: Team Cost, Members, Engagement Rate, Avg Sessions/User. Member table (avatar, name, role, sessions, tokens, cost, last active)."],
        ["5.5", "Cost User Profile", "WS-B", "User-level: profile card, Sessions, Tokens, Audio Minutes, Cost/Session. Cost breakdown bars (Token Costs, Audio Processing, PRISM Assessment, Embeddings)."],
        ["5.6", "Cost Drill-Down Navigation", "WS-B", "Breadcrumb drill-down: Platform > Company > Team > User. Click org > company view, click dept > team view, click user > user profile."],
        ["5.7", "PRISM Assessment Page", "WS-B", "/assessment: take PRISM assessment (questionnaire wizard). Uses PRISM API from Phase 2."],
        ["5.8", "Bedrock Guardrails Integration", "WS-E", "If POC positive from Phase 3: integrate Bedrock Guardrails as middleware layer for LLM input/output filtering. Content safety without changing LLM provider."],
        ["5.9", "Push Notifications", "WS-D", "Service worker push registration. Backend Lambda sends push for: new messages, follow-up reminders, session reminders, issue resolutions."],
    ],
)

doc.add_heading("9.2 Exit Criteria", level=2)
bullet("Cost dashboard shows real cost data at all 4 drill-down levels")
bullet("PRISM Assessment page allows users to take assessments")
bullet("Push notifications work on Chrome/Edge (Safari limited)")
bullet("Bedrock Guardrails filter LLM responses (if adopted)")

doc.add_heading("9.3 Effort", level=2)
doc.add_paragraph("6 weeks, 2\u20133 developers, $36,000\u2013$108,000")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. PHASE 6: INFRASTRUCTURE MODERNIZATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Phase 6: Infrastructure Modernization (Weeks 35\u201344)", level=1)

bold_para("Goal: Extract backend microservices using Strangler Fig pattern.", 12, RGBColor(0, 100, 0))

doc.add_paragraph(
    "This phase runs in parallel with ongoing feature work and bug fixes. The API Gateway "
    "proxy from Phase 3 enables gradual extraction without frontend changes."
)

doc.add_heading("10.1 Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Weeks", "Details"],
    [
        ["6.1", "Extract Auth Service", "WS-D", "35\u201338", "Lambda for /v1/login, /signup, /refresh-token, /me, /social-auth/*, /change-password. Route via API Gateway. Cognito JWT authorizer."],
        ["6.2", "Extract Document Service", "WS-D", "35\u201337", "Lambda for /v1/file_service/*. S3 for blobs, DynamoDB for metadata. Self-contained."],
        ["6.3", "Extract Audit Service", "WS-D", "37\u201338", "Lambda for /v1/audit/*. DynamoDB with TTL. Fire-and-forget pattern preserved."],
        ["6.4", "EventBridge Setup", "WS-D", "38\u201339", "Create inspire-genius-events custom bus. Define event schemas. Wire initial events: user.created, auth.login, document.uploaded."],
        ["6.5", "Extract Chat Service", "WS-D", "39\u201342", "Lambda for /v1/chat/*. WebSocket API Gateway for agent connections. DynamoDB for conversations/messages. Most complex extraction."],
        ["6.6", "Extract Remaining Services", "WS-D", "42\u201344", "Lambda for: Coach, Feedback, User Management, Organization, Prompt, Dashboard. Each follows established pattern."],
        ["6.7", "Monitoring Stack", "WS-D", "43\u201344", "CloudWatch dashboards per service. X-Ray distributed tracing. Alarms for error rates and latency. Cost anomaly alerts."],
    ],
)

doc.add_heading("10.2 Extraction Order Rationale", level=2)
bullet("Auth first: clear boundaries, foundation for JWT authorizer on all other services")
bullet("Document + Audit next: self-contained, minimal dependencies, low risk")
bullet("Chat last (among initial batch): most complex due to WebSocket, audio streaming")
bullet("Remaining services: follow established pattern, faster extraction")

doc.add_heading("10.3 Frontend Impact", level=2)
doc.add_paragraph(
    "Zero frontend changes required. The API Gateway preserves the same /v1/* path structure. "
    "The only change is VITE_API_BASE_URL pointing to the API Gateway domain instead of the "
    "monolith directly. WebSocket URLs change to API Gateway WebSocket endpoints."
)

doc.add_heading("10.4 Exit Criteria", level=2)
bullet("All 12 services running as Lambda behind API Gateway")
bullet("Monolith receives zero traffic (can be decommissioned)")
bullet("EventBridge events flowing between services")
bullet("CloudWatch dashboards operational for all services")
bullet("No user-facing regression (latency within SLO)")

doc.add_heading("10.5 Effort", level=2)
doc.add_paragraph("10 weeks, 2 developers, $60,000\u2013$120,000")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. PHASE 7: SUPER ADMIN & POLISH
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Phase 7: Super Admin & Polish (Weeks 45\u201348)", level=1)

bold_para("Goal: Migrate Super Admin into unified layout. Final polish and optimization.", 12, RGBColor(0, 100, 0))

doc.add_heading("11.1 Deliverables", level=2)
add_table(
    ["#", "Deliverable", "Workstream", "Details"],
    [
        ["7.1", "Super Admin Nav Group", "WS-B", "Add ADMIN section to unified sidebar. Migrate: Dashboard, User Mgmt, Coach Mgmt, RLHF, Prompt Builder, Audit Log, Org Mgmt, Settings."],
        ["7.2", "Route Migration", "WS-B", "Move /super-admin/* routes to /admin/*. Add redirects for old URLs. Update all internal links."],
        ["7.3", "Cost Dashboard Access", "WS-B", "Grant Super Admin access to Platform-level Cost Dashboard. Add cost nav item to ADMIN group."],
        ["7.4", "Feature Flags Cleanup", "WS-B", "Remove any feature flags used during transition. Remove old UserLayout and SuperAdminLayout files."],
        ["7.5", "Performance Optimization", "WS-D", "Code-split by route (React.lazy). Preload critical routes. Optimize bundle size. Lighthouse audit > 90."],
        ["7.6", "Accessibility Audit", "WS-B", "WCAG 2.1 AA compliance check. Fix keyboard navigation, screen reader, color contrast issues."],
        ["7.7", "Monolith Decommission", "WS-D", "After 4-week parallel run: shut down monolith EC2/ECS instances. Update DNS. Archive old code."],
        ["7.8", "Documentation Update", "WS-B", "Update CLAUDE.md, change_log.md, database_schema.md, IG_project_log.html with all changes."],
    ],
)

doc.add_heading("11.2 Exit Criteria", level=2)
bullet("Super Admin section works in unified layout with all existing functionality preserved")
bullet("Old /super-admin/* routes redirect correctly")
bullet("Lighthouse PWA score > 90, Performance > 80")
bullet("Monolith fully decommissioned")
bullet("All documentation up to date")

doc.add_heading("11.3 Effort", level=2)
doc.add_paragraph("4 weeks, 1\u20132 developers, $12,000\u2013$48,000")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. CONSOLIDATED TIMELINE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Consolidated Timeline", level=1)

add_table(
    ["Phase", "Weeks", "Calendar (from start)", "Primary Focus", "Workstreams Active"],
    [
        ["0: Stabilize", "1\u20134", "Month 1", "Fix critical bugs, document environment", "WS-A, WS-E"],
        ["1: Foundation", "5\u20138", "Month 2", "Multi-role auth, unified layout, PWA manifest, Secrets Manager", "WS-A, WS-B, WS-C, WS-D, WS-E"],
        ["2: Core User", "9\u201314", "Months 3\u20134", "User dashboard, PRISM, Goals, Service Worker, Notifications", "WS-A, WS-B, WS-C, WS-D"],
        ["3: Mgmt & Org", "15\u201322", "Months 4\u20136", "Manager (10 pages), Company Admin (6 pages), CDK, API Gateway proxy, Bedrock POC", "WS-A, WS-B, WS-C, WS-D, WS-E"],
        ["4: Pract & Dist", "23\u201328", "Months 6\u20137", "Practitioner (5 pages), Distributor (4 pages), Credits, Messages", "WS-A, WS-B, WS-C"],
        ["5: Intelligence", "29\u201334", "Months 8\u20139", "Cost Dashboard (4 levels), Assessment, Guardrails, Push Notifications", "WS-A, WS-B, WS-C, WS-D, WS-E"],
        ["6: Infra Modern", "35\u201344", "Months 9\u201311", "Microservices extraction (12 services), EventBridge, Monitoring", "WS-A, WS-D"],
        ["7: SA & Polish", "45\u201348", "Month 12", "Super Admin migration, performance, accessibility, monolith decommission", "WS-A, WS-B, WS-D"],
    ],
)

doc.add_paragraph()
bold_para("Total Duration: 48 weeks (12 months)", 13, RGBColor(0, 51, 102))

doc.add_heading("12.1 Accelerated Timeline (Team of 3+)", level=2)
doc.add_paragraph(
    "With 3 developers working in parallel, phases 2\u20135 can overlap significantly. "
    "Accelerated estimate: 32\u201336 weeks (8\u20139 months)."
)

doc.add_heading("12.2 Minimum Viable Timeline (Solo Developer)", level=2)
doc.add_paragraph(
    "Solo developer working full-time: 14\u201318 months. Recommend deferring Phase 6 "
    "(microservices) and Phase 7 (Super Admin migration) until additional resources are available."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 13. RESOURCE REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Resource Requirements", level=1)

doc.add_heading("13.1 Team Composition", level=2)
add_table(
    ["Role", "Count", "Skills Required", "Phases Active"],
    [
        ["Senior Frontend Dev", "1\u20132", "React 19, TypeScript, Tailwind, shadcn/ui, Recharts, Vite, PWA/Workbox", "All phases"],
        ["Senior Backend Dev", "1", "Python/FastAPI (current stack), AWS Lambda, DynamoDB, API Gateway, EventBridge", "Phases 1\u20136"],
        ["DevOps / Infra", "0.5 (part-time)", "AWS CDK, CloudFormation, CI/CD, monitoring, CloudWatch, X-Ray", "Phases 3, 5, 6, 7"],
        ["UI/UX Designer", "0.25 (advisory)", "Figma, responsive design, accessibility, component design system", "Phases 1\u20134"],
        ["QA / Testing", "0.5 (part-time)", "E2E testing (Playwright/Cypress), API testing, cross-browser, mobile PWA", "Phases 2\u20137"],
    ],
)

doc.add_heading("13.2 Cost Summary", level=2)
add_table(
    ["Phase", "Duration", "Team", "Cost Range"],
    [
        ["0: Stabilize", "4 weeks", "1\u20132 devs", "$6,000 \u2013 $24,000"],
        ["1: Foundation", "4 weeks", "2 devs", "$24,000 \u2013 $48,000"],
        ["2: Core User", "6 weeks", "2\u20133 devs", "$36,000 \u2013 $108,000"],
        ["3: Mgmt & Org", "8 weeks", "2\u20133 devs", "$48,000 \u2013 $144,000"],
        ["4: Pract & Dist", "6 weeks", "2\u20133 devs", "$36,000 \u2013 $108,000"],
        ["5: Intelligence", "6 weeks", "2\u20133 devs", "$36,000 \u2013 $108,000"],
        ["6: Infra Modern", "10 weeks", "2 devs", "$60,000 \u2013 $120,000"],
        ["7: SA & Polish", "4 weeks", "1\u20132 devs", "$12,000 \u2013 $48,000"],
        ["TOTAL", "48 weeks", "\u2014", "$258,000 \u2013 $708,000"],
    ],
)

doc.add_paragraph()
bold_para("Solo developer estimate: $108,000 \u2013 $270,000 over 14\u201318 months", 12, RGBColor(0, 51, 102))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 14. GO/NO-GO DECISION GATES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Go/No-Go Decision Gates", level=1)

doc.add_paragraph(
    "Each phase boundary includes a decision gate. The team reviews progress and decides "
    "whether to proceed, adjust scope, or pause."
)

add_table(
    ["Gate", "Decision Point", "Go Criteria", "No-Go Action"],
    [
        ["Gate 0\u21921", "End of Phase 0", "Core flows work (login, chat, documents). No Critical bugs remain.", "Continue Phase 0 until stable."],
        ["Gate 1\u21922", "End of Phase 1", "Unified layout works. Multi-role auth tested. No regressions.", "Fix layout/auth issues before proceeding."],
        ["Gate 2\u21923", "End of Phase 2", "User section complete. PWA installs. PRISM renders.", "Fix PRISM API issues. PWA can be deferred."],
        ["Gate 3\u21924", "End of Phase 3", "Manager + Company sections work. API Gateway proxy stable.", "Reduce scope: defer Hiring/Job DNA. Fix API Gateway."],
        ["Gate 4\u21925", "End of Phase 4", "Practitioner + Distributor work. Credit system accurate.", "Fix credit ledger. Cannot proceed with inaccurate financials."],
        ["Gate 5\u21926", "End of Phase 5", "Cost dashboard shows real data. Guardrails decision made.", "Cost dashboard can launch with mock data initially."],
        ["Gate 6\u21927", "End of Phase 6", "All microservices running. Monolith traffic at zero. No latency regression.", "Keep monolith running. Fix failing services."],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 15. RISK REGISTER
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("15. Risk Register", level=1)

add_table(
    ["ID", "Risk", "Phase", "Likelihood", "Impact", "Mitigation"],
    [
        ["R1", "Backend APIs not ready when frontend needs them", "2\u20135", "High", "High", "Frontend builds with MSW (Mock Service Worker). Integrate real APIs when ready."],
        ["R2", "Multi-role auth introduces security vulnerabilities", "1", "Medium", "Critical", "Security review of hasAccess(). Unit test all role combinations. Penetration test."],
        ["R3", "PRISM assessment data not available from backend", "2", "Medium", "High", "Build PRISM components with sample data. Implement real API as a follow-up."],
        ["R4", "Credit system financial inaccuracy", "4", "Medium", "Critical", "Server-side ledger is source of truth. Frontend is read-only. Double-entry accounting pattern."],
        ["R5", "Org chart performance at scale (1000+ nodes)", "3", "Medium", "Medium", "Virtual rendering, lazy-load subtrees, limit initial depth to 3 levels."],
        ["R6", "Scope creep across 35+ new routes", "2\u20135", "High", "High", "Ship one section per phase. Each phase independently deployable. Cut scope before slipping timeline."],
        ["R7", "WebSocket migration breaks real-time chat", "6", "High", "Critical", "Extensive testing. Canary deployment (10% traffic to new, 90% to old). Instant rollback capability."],
        ["R8", "PWA service worker caches stale data", "2", "Medium", "Medium", "autoUpdate strategy. Version-based cache busting. Clear cache on deploy."],
        ["R9", "Team burnout across 12-month plan", "All", "Medium", "High", "Ship wins early (each phase has visible value). Celebrate milestones. Adjust pace if needed."],
        ["R10", "Breaking changes to existing user workflows", "1\u20132", "Medium", "High", "Feature flag new UI. Opt-in period. Run old + new in parallel. Gather user feedback."],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 16. SUCCESS METRICS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("16. Success Metrics", level=1)

doc.add_heading("16.1 Per-Phase Metrics", level=2)
add_table(
    ["Phase", "Key Metric", "Target"],
    [
        ["0: Stabilize", "Critical bug count", "0 critical bugs remaining"],
        ["1: Foundation", "Multi-role auth accuracy", "100% correct access for all role combinations"],
        ["2: Core User", "PWA Lighthouse score", "> 90 (PWA category)"],
        ["2: Core User", "PRISM report load time", "< 2 seconds"],
        ["3: Mgmt & Org", "Manager feature adoption", "> 50% of managers using within 30 days"],
        ["4: Pract & Dist", "Credit ledger accuracy", "100% (zero discrepancies)"],
        ["5: Intelligence", "Cost dashboard data freshness", "< 1 hour delay from real-time"],
        ["6: Infra Modern", "P95 API latency", "< 500ms (same or better than monolith)"],
        ["6: Infra Modern", "Deployment frequency", "< 10 min deploy per service (vs 30+ min monolith)"],
        ["7: SA & Polish", "Lighthouse Performance score", "> 80"],
        ["7: SA & Polish", "Bundle size", "< 500KB initial load (code-split)"],
    ],
)

doc.add_heading("16.2 Overall Program Metrics", level=2)
add_table(
    ["Metric", "Current Baseline", "12-Month Target"],
    [
        ["Number of user roles supported", "5 (but only 2 functional: user, super-admin)", "6 fully functional roles"],
        ["Total application routes", "~25", "~65"],
        ["PWA installability", "No", "Yes (all platforms)"],
        ["Offline capability", "None", "App shell + cached data"],
        ["Backend services", "1 monolith + 2 satellites", "12 independent microservices"],
        ["Deployment time", "5+ minutes (full rebuild)", "< 1 minute per service"],
        ["API key management", "3 keys in env vars", "0 keys (IAM + Secrets Manager)"],
        ["Infrastructure as Code", "0%", "100% (AWS CDK)"],
        ["Time to onboard new developer", "2\u20133 days", "< 1 day (clear service boundaries, docs)"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("\u2014 End of Document \u2014")
run.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── SAVE ────────────────────────────────────────────────────────────────
output_dir = "/Users/williambrown/Dropbox/AES Material/Inspire-X/New IG Projects/Local_IG-App_UI/inspire-genius-frontend"
output_path = os.path.join(output_dir, "Inspire_Genius_Phased_Implementation_Plan.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
