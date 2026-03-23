"""Generate the Inspire Genius Architecture Plan Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0, 51, 102)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


# ── TITLE PAGE ──────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Inspire Genius\nArchitecture Rearchitecture Plan")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    "PWA \u2022 Microservices \u2022 Event-Driven \u2022 Serverless Lambda"
)
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0, 128, 128)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Prepared for: ").bold = True
meta.add_run("AES / Inspire Genius Team\n")
meta.add_run("Date: ").bold = True
meta.add_run("March 13, 2026\n")
meta.add_run("Classification: ").bold = True
meta.add_run("Internal \u2014 Strategic Planning")

doc.add_page_break()

# ── TABLE OF CONTENTS ───────────────────────────────────────────────────
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Current State Assessment",
    "3. PWA Conversion Plan",
    "4. Microservices Decomposition",
    "5. Event-Driven Architecture",
    "6. Serverless Lambda Architecture",
    "7. Infrastructure as Code",
    "8. Pros and Cons",
    "9. Benefits and Risks",
    "10. Costs and Expenses",
    "11. Timing Analysis & Recommendation",
    "12. Appendix: Complete API Endpoint Map",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Number")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "Inspire Genius is an AI-powered coaching platform that currently operates as a "
    "monolithic Single Page Application (SPA) frontend deployed to AWS S3 + CloudFront, "
    "communicating with a monolithic backend API and two smaller satellite services "
    "(Audit and Magic Auth). The platform serves coaches, users, and administrators "
    "across organizations with features including real-time AI chat (text + voice), "
    "document management, RLHF feedback collection, and multi-tenant organization management."
)
doc.add_paragraph(
    "This document proposes rearchitecting the platform into four complementary paradigms:"
)
bullet("Progressive Web App (PWA) \u2014 offline capability, installability, push notifications")
bullet("Microservices \u2014 12 bounded-context services replacing the monolith backend")
bullet("Event-Driven Architecture \u2014 Amazon EventBridge for asynchronous inter-service communication")
bullet("Serverless Lambda \u2014 AWS Lambda functions behind API Gateway, eliminating server management")

doc.add_paragraph(
    "The estimated development cost ranges from $120,000\u2013$378,000 (team) or $72,000\u2013$180,000 "
    "(solo developer) over 22\u201331 weeks. Monthly AWS infrastructure costs are projected at "
    "$175\u2013$538, comparable to the current estimated $130\u2013$480/month. The recommended approach "
    "is a phased migration using the Strangler Fig pattern, starting after critical bugs from "
    "the recent server migration are stabilized."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. CURRENT STATE ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Current State Assessment", level=1)

doc.add_heading("2.1 Frontend Architecture", level=2)
add_table(
    ["Component", "Technology", "Details"],
    [
        ["Framework", "React 19 + TypeScript", "Functional components, strict TS"],
        ["Build Tool", "Vite 7", "Fast HMR, esbuild bundling"],
        ["Styling", "Tailwind CSS 4 + shadcn/ui", "Utility-first, Radix primitives"],
        ["State (Server)", "TanStack React Query", "49 service files, 60 hooks"],
        ["State (Client)", "React Context", "AuthContext, TourContext, SidebarContext"],
        ["Routing", "React Router v6", "~25 routes, ProtectedRoute guard"],
        ["Forms", "React Hook Form + Zod", "Schema validation"],
        ["Deployment", "S3 + CloudFront", "GitLab CI \u2192 build \u2192 deploy"],
    ],
)

doc.add_heading("2.2 API Layer", level=2)
doc.add_paragraph("The frontend communicates through three separate axios instances:")
add_table(
    ["Instance", "Base URL Env Var", "Auth Method", "Purpose"],
    [
        ["Main API", "VITE_API_BASE_URL", "access-token header + 401 refresh", "All core endpoints (/v1/*)"],
        ["Audit API", "VITE_AUDIT_SERVICE_URL", "X-API-Key header", "Audit logging (fire-and-forget)"],
        ["Magic Auth API", "VITE_MAGIC_AUTH_URL", "Bearer token", "Magic link / OTP auth"],
    ],
)

doc.add_heading("2.3 WebSocket Connections", level=2)
add_table(
    ["Connection", "Env Var", "Protocol", "Purpose"],
    [
        ["Agent Chat", "VITE_AGENTS_WEBSOCKET_BASE_URL", "JSON + Binary (audio)", "Real-time AI coach conversations"],
        ["Alex Voice", "VITE_ALEX_WEB_SOCKET_URL", "JSON + Binary (audio)", "Voice assistant"],
    ],
)

doc.add_heading("2.4 What Does NOT Exist", level=2)
bullet("No service worker or PWA manifest")
bullet("No offline capability")
bullet("No push notifications")
bullet("No Infrastructure-as-Code (no CDK, CloudFormation, Terraform, or SAM)")
bullet("No backend code in this repository (backend is a separate monolith)")

doc.add_heading("2.5 Current Deployment Pipeline", level=2)
doc.add_paragraph(
    "GitLab CI pipeline (development branch only): npm ci \u2192 npm run build \u2192 "
    "SonarQube scan \u2192 aws s3 sync to S3 bucket \u2192 CloudFront cache invalidation."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. PWA CONVERSION PLAN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. PWA Conversion Plan", level=1)

doc.add_heading("3.1 Web App Manifest", level=2)
doc.add_paragraph(
    "Create public/manifest.webmanifest with app name, icons (192px, 512px, maskable), "
    "start_url: /home, display: standalone, theme_color matching the app\u2019s dark blue palette. "
    "Add <link rel=\"manifest\"> to index.html along with Apple-specific meta tags for iOS PWA support."
)

doc.add_heading("3.2 Service Worker Strategy", level=2)
doc.add_paragraph(
    "Use vite-plugin-pwa (Workbox integration for Vite) with autoUpdate registration type."
)
add_table(
    ["Resource Type", "Caching Strategy", "Rationale"],
    [
        ["App shell (HTML)", "NetworkFirst, fallback to cache", "Always try fresh, work offline"],
        ["JS/CSS bundles", "CacheFirst (hashed filenames)", "Immutable content-addressed files"],
        ["Static images/fonts", "CacheFirst, 30-day expiry", "Rarely change"],
        ["API GET responses", "NetworkFirst, 5-min cache", "Show stale data when offline"],
        ["API mutations (POST/PUT/DELETE)", "NetworkOnly + Background Sync", "Queue failed writes for retry"],
        ["WebSocket", "No caching", "Real-time only, degrade gracefully"],
        ["Document uploads/downloads", "NetworkOnly", "Too large and variable to cache"],
    ],
)

doc.add_heading("3.3 Offline-First Capabilities", level=2)
doc.add_paragraph("Tiered offline support:")
bullet("Tier 1 (Immediate): App shell loads offline, \u201cYou are offline\u201d banner, cached coach list and conversation history via React Query persisted cache (IndexedDB)")
bullet("Tier 2 (Phase 2): Background Sync for queued mutations (feedback, help tickets), offline document viewer for recently viewed PDFs")
bullet("Tier 3 (Future): Full offline chat composition with message queuing, offline onboarding form completion")

doc.add_heading("3.4 Push Notifications", level=2)
doc.add_paragraph(
    "Register PushManager subscription in service worker, POST subscription endpoint to backend. "
    "Notification types: new coach message, document shared, issue status update, session reminder. "
    "Backend Lambda sends push via Web Push protocol, triggered by EventBridge events."
)

doc.add_heading("3.5 Install Prompt", level=2)
doc.add_paragraph(
    "Create useInstallPrompt hook: listen for beforeinstallprompt event, show install banner "
    "on Home page after 3rd visit, track installation via appinstalled event for analytics."
)

doc.add_heading("3.6 Files to Create/Modify", level=2)
add_table(
    ["File", "Change"],
    [
        ["vite.config.ts", "Add VitePWA plugin configuration"],
        ["index.html", "Add manifest link, theme-color meta, apple-touch-icon"],
        ["src/App.tsx", "Add OfflineBanner component, install prompt"],
        ["src/main.tsx", "Register service worker"],
        ["public/manifest.webmanifest", "New \u2014 PWA manifest"],
        ["public/images/", "Add PWA icons (192, 512, maskable)"],
        ["src/hooks/useInstallPrompt.ts", "New \u2014 install prompt hook"],
        ["src/hooks/useOnlineStatus.ts", "New \u2014 online/offline detection hook"],
        ["src/components/shared/OfflineBanner.tsx", "New \u2014 offline indicator"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. MICROSERVICES DECOMPOSITION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Microservices Decomposition", level=1)

doc.add_heading("4.1 Domain-Driven Service Boundaries", level=2)
doc.add_paragraph(
    "The monolith backend should be decomposed into 12 bounded-context microservices, "
    "each owning its data store and API surface."
)
add_table(
    ["#", "Service", "Endpoints Owned", "Data Store"],
    [
        ["1", "Identity & Auth", "/v1/login, /signup, /refresh-token, /me, /social-auth/*, /change-password", "Cognito + DynamoDB"],
        ["2", "Coach / Agent", "/v1/agents-settings/*, /v1/coaches/*", "DynamoDB"],
        ["3", "Chat & Conversation", "/v1/chat/conversations/*, /v1/chat/sessions/*", "DynamoDB"],
        ["4", "Alex Voice Assistant", "/v1/chat/AlexChat/*", "DynamoDB"],
        ["5", "Document", "/v1/file_service/*", "S3 + DynamoDB"],
        ["6", "Organization & Tenant", "/v1/organizations/*, /v1/licenses/*", "Aurora Serverless v2"],
        ["7", "User Management", "/v1/user-management/*, /v1/rbac/*, /v1/profile", "DynamoDB + Cognito"],
        ["8", "Feedback & RLHF", "/v1/feedback/*", "DynamoDB"],
        ["9", "Support & Issues", "/v1/issues/*", "DynamoDB"],
        ["10", "Prompt Management", "/v1/prompts/*", "DynamoDB"],
        ["11", "Dashboard & Analytics", "/v1/dashboard/*", "DynamoDB (CQRS aggregates)"],
        ["12", "Audit", "/v1/audit/*", "DynamoDB (TTL retention)"],
    ],
)

doc.add_heading("4.2 API Gateway Configuration", level=2)
doc.add_paragraph(
    "Use AWS API Gateway HTTP API (v2) for all synchronous HTTP endpoints. HTTP API is ~70% cheaper "
    "than REST API, has lower latency (~10ms vs ~30ms), and supports JWT authorizers natively via Cognito."
)
bullet("Single custom domain: api.inspiregenius.com")
bullet("Path-based routing to individual Lambda functions")
bullet("Cognito JWT Authorizer on all routes")
bullet("CORS configured per stage")
bullet("Rate limiting: 1000 req/s default, 100 req/s per user")

doc.add_paragraph("Two separate WebSocket API Gateways:")
bullet("ws-agents.inspiregenius.com \u2192 Agent Chat Lambda ($connect, $disconnect, $default routes)")
bullet("ws-alex.inspiregenius.com \u2192 Alex Voice Lambda (same pattern)")

doc.add_heading("4.3 Data Ownership Principle", level=2)
doc.add_paragraph(
    "Each service owns its data store exclusively. Cross-service data access happens via:"
)
bullet("Events (preferred): Service A emits event, Service B updates its local copy")
bullet("Synchronous query (exception): Dashboard service may query other services\u2019 read APIs")
doc.add_paragraph(
    "Example: When a user is created (Auth Service), a user.created event is emitted. "
    "User Management Service and Organization Service consume it to update their local state."
)

doc.add_heading("4.4 Frontend Impact", level=2)
doc.add_paragraph(
    "The frontend changes are minimal. The API Gateway preserves the same /v1/* path structure. "
    "The only change is updating VITE_API_BASE_URL to point to the API Gateway domain instead of "
    "the monolith. The existing Service \u2192 Hook \u2192 Component pattern, axios interceptors, "
    "and WebSocket hooks remain unchanged."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. EVENT-DRIVEN ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Event-Driven Architecture", level=1)

doc.add_heading("5.1 Event Bus: Amazon EventBridge", level=2)
doc.add_paragraph(
    "Use a custom EventBridge bus named inspire-genius-events. EventBridge is preferred over "
    "SNS/SQS because it provides: schema registry for event contracts, content-based filtering, "
    "archive and replay capability (critical for debugging), built-in retry with DLQ, and "
    "pay-per-event pricing ($1/million events)."
)

doc.add_heading("5.2 Event Catalog", level=2)
add_table(
    ["Event", "Source", "Consumers", "Pattern"],
    [
        ["user.created", "Auth Service", "User Mgmt, Org Service, Audit", "Notification"],
        ["user.updated", "User Mgmt", "Audit, Dashboard", "Notification"],
        ["user.deleted", "User Mgmt", "Auth, Org, Audit", "Notification"],
        ["user.invited", "User Mgmt", "Auth (create Cognito user), Audit", "Command"],
        ["auth.login", "Auth Service", "Audit, Dashboard", "Notification"],
        ["auth.logout", "Auth Service", "Audit", "Notification"],
        ["org.created", "Org Service", "Audit, Dashboard", "Notification"],
        ["org.deactivated", "Org Service", "User Mgmt, Audit", "Notification"],
        ["license.created", "Org Service", "Audit", "Notification"],
        ["coach.assigned", "Coach Service", "Org Service, Audit", "Notification"],
        ["conversation.started", "Chat Service", "Audit, Dashboard", "Notification"],
        ["message.sent", "Chat Service", "Audit, Dashboard, Feedback", "Notification"],
        ["document.uploaded", "Document Service", "Audit, Dashboard", "Notification"],
        ["document.deleted", "Document Service", "Audit", "Notification"],
        ["feedback.submitted", "Feedback Service", "Audit, Dashboard, Prompt Mgmt", "Notification"],
        ["issue.created", "Support Service", "Audit, Dashboard", "Notification"],
        ["issue.resolved", "Support Service", "Audit, Push Notification", "Notification"],
        ["prompt.updated", "Prompt Service", "Coach Service, Audit", "Notification"],
    ],
)

doc.add_heading("5.3 Saga Pattern for Distributed Transactions", level=2)
doc.add_paragraph("Saga 1 \u2014 User Invitation Flow:")
bullet("User Mgmt receives invite request \u2192 creates invitation record \u2192 emits user.invited")
bullet("Auth Service creates Cognito user \u2192 emits cognito.user.created")
bullet("User Mgmt updates invitation status \u2192 Notification Service sends email")
bullet("Compensation: If Cognito creation fails, User Mgmt marks invitation as failed")

doc.add_paragraph("Saga 2 \u2014 Organization Onboarding:")
bullet("Org Service creates organization \u2192 creates license \u2192 assigns coaches")
bullet("Compensation: If coach assignment fails, roll back license and org")

doc.add_paragraph("Implementation: AWS Step Functions for orchestrated sagas.")

doc.add_heading("5.4 CQRS Application", level=2)
doc.add_paragraph(
    "Apply CQRS to the Dashboard & Analytics Service: the write side consumes events from all "
    "other services and updates pre-computed aggregates in DynamoDB. The read side is a simple "
    "Lambda that reads from the aggregates table. This avoids the dashboard needing to fan out "
    "queries to 10+ services on every page load."
)
doc.add_paragraph(
    "Also apply to the Audit Service: write side receives events and writes to DynamoDB with TTL; "
    "read side supports paginated queries via GSI on timestamp."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. SERVERLESS LAMBDA ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Serverless Lambda Architecture", level=1)

doc.add_heading("6.1 Lambda Function Design", level=2)
doc.add_paragraph(
    "One Lambda per service (not one per endpoint). Use internal routing via a lightweight "
    "router or manual path matching. Runtime: Node.js 22.x. Bundler: esbuild."
)
add_table(
    ["Lambda", "Memory", "Timeout", "Reserved Concurrency"],
    [
        ["Auth", "256 MB", "10s", "100"],
        ["Coach/Agent", "256 MB", "10s", "50"],
        ["Chat & Conversation", "512 MB", "30s", "200"],
        ["Alex Voice", "512 MB", "30s", "100"],
        ["Document", "512 MB", "60s", "50"],
        ["Organization", "256 MB", "10s", "50"],
        ["User Management", "256 MB", "10s", "50"],
        ["Feedback", "256 MB", "10s", "50"],
        ["Support", "256 MB", "10s", "30"],
        ["Prompt", "256 MB", "10s", "20"],
        ["Dashboard", "256 MB", "10s", "50"],
        ["Audit", "256 MB", "10s", "100"],
        ["WS Agent Handler", "512 MB", "30s", "200"],
        ["WS Alex Handler", "512 MB", "30s", "100"],
        ["Event Processors (\u00d75)", "256 MB", "30s", "50 each"],
    ],
)

doc.add_heading("6.2 Cold Start Mitigation", level=2)
bullet("Provisioned Concurrency for Auth (5 instances) and Chat (10 instances)")
bullet("Keep bundle sizes small (<5 MB per Lambda)")
bullet("Lazy-load AWS SDK clients")
bullet("Use Lambda SnapStart when available for Node.js")

doc.add_heading("6.3 Data Storage Decisions", level=2)
add_table(
    ["Service", "Storage", "Rationale"],
    [
        ["Auth", "Cognito + DynamoDB", "Cognito handles identity; DynamoDB for session metadata"],
        ["Coach/Agent", "DynamoDB", "Simple key-value lookups, no relational joins"],
        ["Chat", "DynamoDB", "High write throughput, TTL for old messages, GSI on conversation_id"],
        ["Documents", "S3 + DynamoDB", "S3 for file blobs, DynamoDB for metadata/indexing"],
        ["Organizations", "Aurora Serverless v2", "Complex relational data needing JOINs (org \u2192 business \u2192 license \u2192 coach)"],
        ["User Management", "DynamoDB + Cognito", "User records + Cognito for identity"],
        ["Feedback", "DynamoDB", "Append-only records, GSI for stats queries"],
        ["Support", "DynamoDB", "Issue tickets with status tracking"],
        ["Prompts", "DynamoDB", "Versioned documents, GSI on coach_id"],
        ["Dashboard", "DynamoDB", "Pre-computed aggregates from CQRS"],
        ["Audit", "DynamoDB", "High-volume writes, TTL for retention"],
    ],
)

doc.add_heading("6.4 WebSocket API Gateway Architecture", level=2)
doc.add_paragraph(
    "Client \u2194 API Gateway WebSocket \u2194 Lambda Handler \u2194 DynamoDB (connections table) "
    "\u2194 AI/LLM Service. The Lambda handler receives $connect (auth), $disconnect (cleanup), "
    "and $default (message handling) routes. Responses stream back via the API Gateway Management "
    "API\u2019s postToConnection method."
)

doc.add_heading("6.5 S3 Buckets", level=2)
add_table(
    ["Bucket", "Purpose"],
    [
        ["inspire-genius-frontend-{env}", "Static frontend assets (existing)"],
        ["inspire-genius-documents-{env}", "User-uploaded documents"],
        ["inspire-genius-org-assets-{env}", "Organization logos"],
        ["inspire-genius-chat-exports-{env}", "Generated chat export PDFs/CSVs"],
    ],
)

doc.add_heading("6.6 Cognito Enhancements", level=2)
bullet("Custom attributes: custom:role, custom:organization_id, custom:business_id")
bullet("Pre-token-generation Lambda trigger to inject role claims into JWT")
bullet("Post-confirmation Lambda trigger to emit user.created event to EventBridge")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. INFRASTRUCTURE AS CODE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Infrastructure as Code", level=1)

doc.add_heading("7.1 AWS CDK (TypeScript)", level=2)
doc.add_paragraph(
    "AWS CDK is recommended over SAM or CloudFormation because: TypeScript aligns with the "
    "existing codebase, it is more expressive than YAML/JSON, L2/L3 constructs reduce boilerplate, "
    "and types can be shared between Lambda code and infrastructure code."
)

doc.add_heading("7.2 CDK Project Structure", level=2)
doc.add_paragraph(
    "infrastructure/\n"
    "  bin/app.ts                     \u2014 CDK app entry point\n"
    "  lib/stacks/\n"
    "    network-stack.ts             \u2014 VPC (for Aurora), security groups\n"
    "    auth-stack.ts                \u2014 Cognito, Auth Lambda, API Gateway routes\n"
    "    api-gateway-stack.ts         \u2014 HTTP API, custom domain, authorizer\n"
    "    ws-gateway-stack.ts          \u2014 WebSocket APIs (agent + alex)\n"
    "    chat-stack.ts                \u2014 Chat Lambda, DynamoDB\n"
    "    document-stack.ts            \u2014 Document Lambda, S3 buckets\n"
    "    org-stack.ts                 \u2014 Organization Lambda, Aurora Serverless\n"
    "    user-mgmt-stack.ts           \u2014 User Management Lambda, DynamoDB\n"
    "    coach-stack.ts               \u2014 Coach Lambda, DynamoDB\n"
    "    feedback-stack.ts            \u2014 Feedback Lambda, DynamoDB\n"
    "    support-stack.ts             \u2014 Support Lambda, DynamoDB\n"
    "    prompt-stack.ts              \u2014 Prompt Lambda, DynamoDB\n"
    "    dashboard-stack.ts           \u2014 Dashboard Lambda, DynamoDB\n"
    "    audit-stack.ts               \u2014 Audit Lambda, DynamoDB\n"
    "    event-bus-stack.ts           \u2014 EventBridge bus, rules, DLQs\n"
    "    frontend-stack.ts            \u2014 S3 bucket, CloudFront, OAI\n"
    "    monitoring-stack.ts          \u2014 CloudWatch dashboards, alarms, X-Ray\n"
    "  lib/constructs/\n"
    "    lambda-service.ts            \u2014 Reusable L3 construct for Lambda + API route\n"
    "    dynamo-table.ts              \u2014 Reusable DynamoDB table construct\n"
    "  config/ dev.ts, staging.ts, prod.ts"
)

doc.add_heading("7.3 CI/CD Pipeline Update", level=2)
doc.add_paragraph("Updated GitLab CI pipeline:")
bullet("Lint + Test (frontend)")
bullet("Build frontend (Vite)")
bullet("CDK synth (validate IaC)")
bullet("CDK diff (show changes)")
bullet("Manual approval gate (for production)")
bullet("CDK deploy (per stack, in dependency order)")
bullet("Deploy frontend to S3 + CloudFront invalidation")
bullet("Run smoke tests")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. PROS AND CONS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Pros and Cons", level=1)

doc.add_heading("8.1 Pros", level=2)
add_table(
    ["#", "Pro", "Details"],
    [
        ["1", "Independent Scaling", "Chat service scales to 200 concurrent Lambdas while Prompt service stays at 20, matching actual load patterns."],
        ["2", "Independent Deployment", "Deploy a fix to Feedback without touching Auth or Chat. No full redeployment required."],
        ["3", "Cost Efficiency", "Pay-per-invocation. Low-traffic services (Prompt, Support) cost nearly nothing when idle."],
        ["4", "Fault Isolation", "If Document service crashes, users can still chat, view dashboard, and submit feedback."],
        ["5", "Team Autonomy", "Different developers can own different services with clear boundaries."],
        ["6", "Technology Flexibility", "Could rewrite Chat service in Python for AI/ML without affecting the rest."],
        ["7", "PWA Benefits", "Offline capability, installability, push notifications improve mobile UX for coaching."],
        ["8", "Operational Simplicity", "No servers to patch, no OS upgrades, no capacity planning for compute."],
    ],
)

doc.add_heading("8.2 Cons", level=2)
add_table(
    ["#", "Con", "Details"],
    [
        ["1", "Distributed Complexity", "Debugging a request spanning Auth \u2192 Chat \u2192 Audit requires distributed tracing (X-Ray)."],
        ["2", "Eventual Consistency", "When an org is created, dashboard won\u2019t show it for 1\u20132 seconds until event propagates."],
        ["3", "Cold Starts", "First invocation after idle adds 200\u2013500ms latency. Provisioned concurrency mitigates but adds cost."],
        ["4", "Local Dev Difficulty", "Running 12+ services locally is harder than one monolith. Need SAM local or LocalStack."],
        ["5", "Testing Complexity", "Integration tests must mock or deploy multiple services. E2E tests become critical."],
        ["6", "Vendor Lock-in", "Deep AWS integration (Lambda, DynamoDB, EventBridge, Cognito) makes cloud migration expensive."],
        ["7", "Monitoring Overhead", "12+ services means 12+ sets of logs, metrics, and alarms to manage."],
        ["8", "Network Latency", "Inter-service communication adds network hops. Event propagation adds latency."],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. BENEFITS AND RISKS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Benefits and Risks", level=1)

doc.add_heading("9.1 Business Benefits", level=2)
bullet("Business Agility \u2014 Ship features faster once services are established. New coaching features don\u2019t risk breaking auth or billing.")
bullet("Faster Feature Delivery \u2014 Smaller codebases, faster builds (30 seconds vs 5+ minutes), faster deploys.")
bullet("Better Fault Isolation \u2014 A memory leak in document processing doesn\u2019t take down the entire platform.")
bullet("Pay-Per-Use Pricing \u2014 Dramatic cost savings during low-traffic periods (nights, weekends). Coaching platforms have strong diurnal patterns.")
bullet("PWA Installability \u2014 Users \u201cinstall\u201d the coaching app on phone/desktop, increasing engagement and retention.")
bullet("Push Notifications \u2014 Proactively nudge users about coaching sessions, documents, issue resolutions \u2014 core to a coaching platform\u2019s value.")
bullet("Offline Access \u2014 Coaches and users in areas with poor connectivity can review past conversations and documents.")
bullet("Compliance & Audit \u2014 Clean service boundaries make it easier to implement data residency, GDPR deletion, and audit trails.")

doc.add_heading("9.2 Risk Matrix", level=2)
add_table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Data loss during migration", "Medium", "Critical", "Blue-green deployment, parallel running, data validation scripts"],
        ["Extended downtime", "Medium", "High", "Phased migration with API Gateway routing; never big-bang"],
        ["Team skill gaps", "High", "High", "Training on AWS serverless, pair programming, start with simplest service"],
        ["Over-engineering", "High", "Medium", "Start with 5\u20136 services, not 12. Merge small services if needed"],
        ["Distributed transaction failures", "Medium", "High", "Sagas with compensating transactions, DLQs, monitoring"],
        ["Cost overruns", "Medium", "Medium", "Billing alarms, monthly architecture review, avoid unnecessary Provisioned Concurrency"],
        ["WebSocket complexity", "High", "High", "API Gateway WebSocket has limits (128KB messages, 2hr idle timeout). Plan fallbacks."],
        ["Vendor lock-in", "Certain", "Medium", "Accept it. AWS is stable. Use abstraction layers for data access."],
        ["Performance regression", "Medium", "Medium", "Benchmark current latencies, set SLOs, monitor with X-Ray"],
        ["Frontend migration complexity", "Low", "Low", "Frontend changes are minimal (just base URL + PWA files)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. COSTS AND EXPENSES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Costs and Expenses", level=1)

doc.add_heading("10.1 Development Costs", level=2)
add_table(
    ["Phase", "Duration", "Team Size", "Cost Range ($75\u2013$150/hr)"],
    [
        ["Phase 0: Stabilize current bugs", "2\u20134 weeks", "1\u20132 devs", "$6,000 \u2013 $24,000"],
        ["Phase 1: PWA conversion", "2\u20133 weeks", "1 dev", "$6,000 \u2013 $18,000"],
        ["Phase 2: IaC foundation + first 3 services", "6\u20138 weeks", "2 devs", "$36,000 \u2013 $96,000"],
        ["Phase 3: Remaining services + event bus", "6\u20138 weeks", "2\u20133 devs", "$36,000 \u2013 $144,000"],
        ["Phase 4: WebSocket migration", "3\u20134 weeks", "2 devs", "$18,000 \u2013 $48,000"],
        ["Phase 5: Testing, monitoring, cutover", "3\u20134 weeks", "2 devs", "$18,000 \u2013 $48,000"],
        ["TOTAL", "22\u201331 weeks", "\u2014", "$120,000 \u2013 $378,000"],
    ],
)
doc.add_paragraph("Solo developer estimate: 8\u201312 months, $72,000\u2013$180,000", style="Intense Quote")

doc.add_heading("10.2 AWS Monthly Infrastructure Costs (Production)", level=2)
doc.add_paragraph("Assumes ~1,000 daily active users, ~50,000 API requests/day, ~500 WebSocket connections/day:")
add_table(
    ["AWS Service", "Monthly Cost", "Notes"],
    [
        ["Lambda", "$15 \u2013 $40", "~1.5M invocations/month, 256MB avg, 200ms avg duration"],
        ["API Gateway HTTP", "$5 \u2013 $15", "$1/million requests"],
        ["API Gateway WebSocket", "$10 \u2013 $30", "Connection minutes + messages"],
        ["DynamoDB (on-demand)", "$25 \u2013 $75", "~10 tables, on-demand capacity"],
        ["Aurora Serverless v2", "$50 \u2013 $150", "0.5\u20134 ACU, scales to zero"],
        ["S3", "$5 \u2013 $15", "Frontend + documents + exports"],
        ["CloudFront", "$10 \u2013 $30", "Frontend CDN + API caching"],
        ["Cognito", "$0 \u2013 $28", "First 50K MAU free, then $0.0055/MAU"],
        ["EventBridge", "$1 \u2013 $5", "$1/million events"],
        ["Step Functions", "$5 \u2013 $15", "$25/million state transitions"],
        ["CloudWatch", "$10 \u2013 $30", "Logs, metrics, alarms"],
        ["X-Ray", "$5 \u2013 $15", "Distributed tracing"],
        ["Route 53", "$2 \u2013 $5", "Hosted zones + queries"],
        ["ACM (SSL)", "$0", "Free SSL certificates"],
        ["Secrets Manager", "$2 \u2013 $5", "API keys, DB credentials"],
        ["Provisioned Concurrency", "$30 \u2013 $80", "Auth (5) + Chat (10) warm instances"],
        ["TOTAL", "$175 \u2013 $538/month", "\u2014"],
    ],
)

doc.add_heading("10.3 Current Cost Comparison", level=2)
add_table(
    ["Item", "Estimated Monthly Cost"],
    [
        ["EC2/ECS for monolith backend", "$50 \u2013 $200"],
        ["S3 + CloudFront (frontend)", "$10 \u2013 $30"],
        ["Database (RDS or similar)", "$50 \u2013 $200"],
        ["Magic Auth Service hosting", "$20 \u2013 $50"],
        ["TOTAL (current)", "$130 \u2013 $480/month"],
    ],
)
doc.add_paragraph(
    "Cost delta: The microservices architecture will cost roughly the same or 10\u201330% more in "
    "AWS infrastructure, but operational overhead savings (no server patching, no capacity planning) "
    "and development velocity gains justify it for a growing platform."
)

doc.add_heading("10.4 Hidden Costs", level=2)
add_table(
    ["Hidden Cost", "Estimate", "Type"],
    [
        ["Learning curve (Lambda, CDK, EventBridge)", "$4,000 \u2013 $12,000", "One-time"],
        ["Operational overhead (monitoring, incident response)", "5\u201310 hours/month", "Ongoing"],
        ["LocalStack Pro (optional local dev)", "$50/month", "Ongoing"],
        ["Staging/testing AWS account", "$50 \u2013 $100/month", "Ongoing"],
        ["Third-party monitoring (Datadog, optional)", "$100 \u2013 $500/month", "Ongoing (CloudWatch sufficient to start)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. TIMING ANALYSIS & RECOMMENDATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Timing Analysis & Recommendation", level=1)

doc.add_heading("11.1 The Question", level=2)
doc.add_paragraph(
    "You are currently fixing bugs after migrating from another server. Should you:\n"
    "A) Fix all bugs first, stabilize, THEN migrate to microservices?\n"
    "B) Start microservices migration now in parallel with bug fixes?\n"
    "C) Take a phased approach?"
)

doc.add_heading("11.2 Recommendation: Option C \u2014 Phased Approach", level=2)

p = doc.add_paragraph()
run = p.add_run("RECOMMENDED: Option C \u2014 Phased Approach with Strangler Fig Pattern")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_heading("Why NOT Option A (fix everything first):", level=3)
bullet("Bug fixing could take weeks or months with no end in sight. New bugs surface constantly.")
bullet("You will be fixing bugs in an architecture you plan to abandon \u2014 wasted effort.")
bullet("However, you cannot ignore bugs because users are affected now.")

doc.add_heading("Why NOT Option B (parallel migration):", level=3)
bullet("Fixing bugs in the monolith while simultaneously building microservices creates split attention and doubled cognitive load.")
bullet("The monolith and new microservices codebases would drift apart, creating integration nightmares.")
bullet("Risk of shipping broken features while attention is on infrastructure.")

doc.add_heading("Why Option C (phased approach):", level=3)
bullet("Addresses immediate user pain (bug fixes) while making incremental architectural progress.")
bullet("The Strangler Fig pattern lets you migrate one service at a time with zero downtime.")
bullet("API Gateway as a transparent proxy is the linchpin: route any endpoint to either the monolith or a new microservice.")
bullet("Each phase delivers standalone value. You can pause between phases if needed.")

doc.add_heading("11.3 Recommended Phase Timeline", level=2)
add_table(
    ["Phase", "Weeks", "Focus", "Key Deliverable"],
    [
        ["0: Stabilize", "1\u20133", "Fix CRITICAL and HIGH-severity bugs only", "Stable baseline: login, chat, documents work reliably"],
        ["1: PWA", "4\u20136", "Frontend-only PWA conversion", "Service worker, manifest, offline banner, install prompt"],
        ["2: IaC Foundation", "7\u20139", "CDK project + API Gateway as transparent proxy", "Codified infrastructure, zero behavior change"],
        ["3: Extract Auth", "10\u201313", "First microservice extraction", "Auth Lambda behind API Gateway (Strangler Fig)"],
        ["4: Extract Document + Audit", "14\u201317", "Self-contained services", "Document Lambda + Audit Lambda"],
        ["5: Extract Chat + WebSocket", "18\u201323", "Most complex extraction", "WebSocket API Gateway + Chat Lambda"],
        ["6: Extract Remaining", "24\u201330", "Org, User Mgmt, Feedback, Support, Prompt, Dashboard", "All services migrated"],
        ["7: Decommission Monolith", "31+", "Monitor 2\u20134 weeks, then shut down", "Monolith fully retired"],
    ],
)

doc.add_heading("11.4 Key Principle: Never Stop Fixing Bugs", level=2)
doc.add_paragraph(
    "Throughout ALL phases, dedicate 20\u201330% of each sprint to bug fixes. The phased approach lets "
    "you ship architectural improvements while maintaining the existing system. Bug fixes on the "
    "monolith are NOT wasted effort during Phases 0\u20132 because the monolith is still serving "
    "production traffic. Starting from Phase 3, as you extract services, bugs in extracted domains "
    "get fixed in the new microservice instead."
)

doc.add_heading("11.5 Decision Framework", level=2)
doc.add_paragraph(
    "Use this framework to decide when to start Phase 1 (PWA):"
)
bullet("Can users log in reliably? \u2192 If NO, stay in Phase 0.")
bullet("Can users chat with coaches? \u2192 If NO, stay in Phase 0.")
bullet("Can users upload/view documents? \u2192 If NO, stay in Phase 0.")
bullet("Are there any data-corruption bugs? \u2192 If YES, stay in Phase 0.")
bullet("All of the above are YES and no data corruption? \u2192 Move to Phase 1.")

doc.add_paragraph(
    "You do NOT need to fix every cosmetic bug, edge case, or minor UI issue before starting "
    "the architectural work. Those can be addressed in parallel during the 20\u201330% bug-fix allocation."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. APPENDIX: COMPLETE API ENDPOINT MAP
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Appendix: Complete API Endpoint Map", level=1)

doc.add_paragraph("This appendix maps every frontend API call to its proposed microservice owner.")

add_table(
    ["Endpoint", "Method", "Current Source", "Target Microservice"],
    [
        ["/v1/login", "POST", "Monolith", "1. Auth Service"],
        ["/v1/signup", "POST", "Monolith", "1. Auth Service"],
        ["/v1/verify-signup", "POST", "Monolith", "1. Auth Service"],
        ["/v1/resend-verification", "POST", "Monolith", "1. Auth Service"],
        ["/v1/refresh-token", "POST", "Monolith", "1. Auth Service"],
        ["/v1/me", "GET", "Monolith", "1. Auth Service"],
        ["/v1/social-auth/login", "GET", "Monolith", "1. Auth Service"],
        ["/v1/request-password-reset", "POST", "Monolith", "1. Auth Service"],
        ["/v1/reset-password", "POST", "Monolith", "1. Auth Service"],
        ["/v1/change-password", "POST", "Monolith", "1. Auth Service"],
        ["/api/auth/request-magic-link", "POST", "Magic Auth", "1. Auth Service (merge)"],
        ["/api/auth/verify-magic-link", "POST", "Magic Auth", "1. Auth Service (merge)"],
        ["/v1/agents-settings/agents", "GET", "Monolith", "2. Coach Service"],
        ["/v1/agents-settings/preferences", "POST/PUT", "Monolith", "2. Coach Service"],
        ["/v1/agents-settings/tone", "GET", "Monolith", "2. Coach Service"],
        ["/v1/chat/conversations", "GET", "Monolith", "3. Chat Service"],
        ["/v1/chat/sessions/start", "POST", "Monolith", "3. Chat Service"],
        ["/v1/chat/conversations/{id}/messages", "GET", "Monolith", "3. Chat Service"],
        ["/v1/chat/AlexChat/*", "GET", "Monolith", "4. Alex Service"],
        ["/v1/file_service/*", "ALL", "Monolith", "5. Document Service"],
        ["/v1/organizations/*", "ALL", "Monolith", "6. Organization Service"],
        ["/v1/licenses/*", "ALL", "Monolith", "6. Organization Service"],
        ["/v1/user-management/*", "ALL", "Monolith", "7. User Management"],
        ["/v1/rbac/roles", "GET", "Monolith", "7. User Management"],
        ["/v1/onboarding/profile", "POST/PUT", "Monolith", "7. User Management"],
        ["/v1/feedback/*", "ALL", "Monolith", "8. Feedback Service"],
        ["/v1/issues/*", "ALL", "Monolith", "9. Support Service"],
        ["/v1/prompts/*", "ALL", "Monolith", "10. Prompt Service"],
        ["/v1/dashboard/*", "GET", "Monolith", "11. Dashboard Service"],
        ["/v1/audit/*", "ALL", "Monolith", "12. Audit Service"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("\u2014 End of Document \u2014")
run.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── SAVE ────────────────────────────────────────────────────────────────
output_dir = "/Users/williambrown/Dropbox/AES Material/Inspire-X/New IG Projects/Local_IG-App_UI/inspire-genius-frontend"
output_path = os.path.join(output_dir, "Inspire_Genius_Architecture_Plan.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")