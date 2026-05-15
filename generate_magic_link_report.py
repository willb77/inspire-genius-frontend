#!/usr/bin/env python3
"""Generate the Magic Link Issues Report Word document."""

import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx not installed. Installing...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(os.path.dirname(SCRIPT_DIR), "Logo-Dark.png")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "Magic_Link_Issues_Report.docx")

doc = Document()

# -- Style adjustments --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    heading_style = doc.styles[f"Heading {level}"]
    heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    heading_style.font.name = "Calibri"

# -- Logo --
if os.path.exists(LOGO_PATH):
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run()
    run.add_picture(LOGO_PATH, width=Inches(2.5))
else:
    print(f"WARNING: Logo not found at {LOGO_PATH}")

# -- Title --
title = doc.add_heading("Inspire Genius — Magic Link Integration:\nIssue Report & Fixes", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# -- Date --
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run("Date: 2026-03-16")
date_run.font.size = Pt(12)
date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
date_run.font.italic = True

doc.add_paragraph()  # spacer

# ============================================================
# 1. Executive Summary
# ============================================================
doc.add_heading("1. Executive Summary", level=1)

p = doc.add_paragraph(
    "After migrating from MFA/OTP to Magic Link authentication, two issues "
    "were identified during user testing:"
)

items = [
    "WebSocket connections for AI coach chat failing due to duplicated URL path",
    "My Documents panel showing empty for authenticated users",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "This report details the root causes, applied fixes, and recommendations "
    "for each issue."
)

# ============================================================
# 2. Issue 1: WebSocket URL Duplication
# ============================================================
doc.add_heading("2. Issue 1: WebSocket URL Duplication", level=1)

# Symptom
doc.add_heading("Symptom", level=2)
doc.add_paragraph(
    "Browser console shows repeated WebSocket connection failures to:"
)
code_para = doc.add_paragraph()
code_run = code_para.add_run(
    "wss://dvw79io0afgrp.cloudfront.net/v1/agents/ws/v1/agents/ws/agents/{agentId}"
)
code_run.font.name = "Consolas"
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_paragraph(
    "Note the doubled /v1/agents/ws path segment. The connection fails "
    "immediately, triggering an infinite reconnection loop."
)

# Root Cause
doc.add_heading("Root Cause", level=2)
doc.add_paragraph(
    "The environment variable VITE_AGENTS_WEBSOCKET_BASE_URL is set to:"
)
env_para = doc.add_paragraph()
env_run = env_para.add_run("wss://dvw79io0afgrp.cloudfront.net/v1/agents/ws")
env_run.font.name = "Consolas"
env_run.font.size = Pt(9)
env_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_paragraph(
    "This base URL already includes the WebSocket path prefix /v1/agents/ws. "
    "The frontend code that constructs the full WebSocket URL appends "
    "/v1/agents/ws/agents/{agentId} to this base, causing the path duplication."
)

doc.add_paragraph("The correct URL should be:")
correct_para = doc.add_paragraph()
correct_run = correct_para.add_run(
    "wss://dvw79io0afgrp.cloudfront.net/v1/agents/ws/agents/{agentId}"
)
correct_run.font.name = "Consolas"
correct_run.font.size = Pt(9)
correct_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

# Fix
doc.add_heading("Fix Applied", level=2)
doc.add_paragraph(
    "Updated the WebSocket URL construction code to append only "
    "/agents/{agentId} to the base URL instead of the full "
    "/v1/agents/ws/agents/{agentId} path."
)

# Impact
doc.add_heading("Impact", level=2)
doc.add_paragraph(
    "All AI coach chat sessions were non-functional — WebSocket connections "
    "failed immediately on every attempt, causing infinite reconnection loops. "
    "This affected 100% of users attempting to use any AI coach."
)

# Summary table
doc.add_heading("Summary", level=2)
table1 = doc.add_table(rows=5, cols=2, style="Light Shading Accent 1")
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

rows_data = [
    ("Field", "Detail"),
    ("Severity", "Critical — all coach chat non-functional"),
    ("Affected Component", "WebSocket connection in AI coach chat"),
    ("Root Cause", "Duplicated /v1/agents/ws path segment in URL construction"),
    ("Resolution", "Append only /agents/{agentId} to base URL"),
]
for i, (key, val) in enumerate(rows_data):
    table1.rows[i].cells[0].text = key
    table1.rows[i].cells[1].text = val
    if i == 0:
        for cell in table1.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

# ============================================================
# 3. Issue 2: Documents Not Loading
# ============================================================
doc.add_heading("3. Issue 2: Documents Not Loading", level=1)

# Symptom
doc.add_heading("Symptom", level=2)
doc.add_paragraph(
    "My Documents panel shows no documents for magic-link authenticated users."
)

# Root Cause
doc.add_heading("Root Cause", level=2)
doc.add_paragraph(
    "Investigation confirmed this is NOT a bug. The documents endpoint "
    "(GET /v1/file_service/list) correctly handles magic auth tokens via "
    "verify_token() and queries by user_id matching the JWT sub claim."
)
doc.add_paragraph(
    "The empty document list is because the test user account has no "
    "uploaded documents. A database query confirmed only 3 users in the "
    "entire system have documents:"
)

doc_users = [
    "charles@prismbrainmapping.com — 1 document",
    "markptully2001@gmail.com — 2 documents",
    "yajnesh.k@pacewisdom.com — 1 document",
]
for item in doc_users:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "The test accounts (willb2@3pp.com, aes@3pp.com, willb77@3pp.com) have "
    "never uploaded documents, so the empty panel is expected behavior."
)

# Fix
doc.add_heading("Resolution", level=2)
doc.add_paragraph(
    "No fix needed — the feature is working correctly. The auth flow for "
    "the documents endpoint was verified: verify_token() handles magic auth "
    "HS256 tokens, extracts the user_id from JWT sub claim, and the file "
    "service queries by that user_id. Magic auth user IDs have been synced "
    "to match IG database user IDs, so ownership lookups are correct."
)

# Impact
doc.add_heading("Impact", level=2)
doc.add_paragraph(
    "No impact — expected behavior. Users who log in with magic link and "
    "upload documents will see them correctly in the My Documents panel."
)

# Summary table
doc.add_heading("Summary", level=2)
table2 = doc.add_table(rows=5, cols=2, style="Light Shading Accent 1")
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

rows_data2 = [
    ("Field", "Detail"),
    ("Severity", "None — expected behavior (not a bug)"),
    ("Affected Component", "My Documents panel / file_service API"),
    ("Root Cause", "Test user account has no uploaded documents"),
    ("Resolution", "No fix needed — feature working correctly"),
]
for i, (key, val) in enumerate(rows_data2):
    table2.rows[i].cells[0].text = key
    table2.rows[i].cells[1].text = val
    if i == 0:
        for cell in table2.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

# ============================================================
# 4. Related: Dual Auth Mode Changes
# ============================================================
doc.add_heading("4. Related: Dual Auth Mode Changes", level=1)

doc.add_paragraph(
    "The following backend changes were made to support magic link authentication "
    "alongside the existing Cognito auth flow:"
)

auth_changes = [
    (
        "users/auth.py — verify_jwt_token()",
        "Now checks for kid header in the JWT: if present, uses Cognito RS256 "
        "verification; if absent, uses Magic Auth HS256 verification with SECRET_KEY."
    ),
    (
        "users/auth.py — verify_token()",
        "Returns user info from JWT claims for magic auth tokens, skipping "
        "the Cognito user attribute lookup."
    ),
    (
        "users/decorators.py — require_admin_role() / require_super_admin_role()",
        "Now checks _auth_source == \"magic_auth\" and uses JWT role claims "
        "directly instead of database RBAC lookup."
    ),
    (
        "Magic Auth Lambda JWT_SECRET",
        "Aligned with IG backend SECRET_KEY to ensure tokens are verifiable "
        "by the backend."
    ),
    (
        "Magic Auth User IDs",
        "Synced to match IG database user IDs so that ownership checks and "
        "data queries return correct results."
    ),
]

for title_text, desc in auth_changes:
    p = doc.add_paragraph()
    bold_run = p.add_run(title_text + " — ")
    bold_run.font.bold = True
    p.add_run(desc)

# ============================================================
# 5. Recommendations
# ============================================================
doc.add_heading("5. Recommendations", level=1)

recommendations = [
    "Remove temporary debug/sync endpoints from Magic Auth Lambda before production deployment.",
    "Consider implementing auto-sync of new user registrations to the magic_auth database to avoid manual syncing.",
    "Monitor WebSocket reconnection behavior after the fix is deployed to confirm stable connections.",
    "Test all coach chat sessions and document operations end-to-end after deployment.",
    "Add integration tests that cover both Cognito and magic-link auth flows for critical endpoints.",
    "Document the dual-auth architecture in the backend README for future developer reference.",
]

for rec in recommendations:
    doc.add_paragraph(rec, style="List Bullet")

# ============================================================
# Footer
# ============================================================
doc.add_paragraph()  # spacer
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run("— End of Report —")
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# -- Save --
doc.save(OUTPUT_PATH)
print(f"Report saved to: {OUTPUT_PATH}")
